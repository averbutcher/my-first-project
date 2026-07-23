import asyncio
import json
import os
import re
import secrets
import sys
import tempfile
import time
from pathlib import Path
from typing import Optional

import anthropic as _anthropic
import bcrypt
import yaml
from fastapi import Cookie, Depends, FastAPI, File, Form, HTTPException, Query, Request, Response, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from yaml.loader import SafeLoader
from dotenv import load_dotenv

sys.path.insert(0, str(Path(__file__).parent))

from analyzer import analyze_tender, distill_knowledge, SYSTEM_PROMPT
from engine import load_config, save_config, parse_message, parse_excel, compare, export_to_excel, find_suspicious_lines
from matching import BranchMatcher, NameMatcher, build_worker_matcher, normalize_match_text
from scraper import Tender, _extract_id_from_url, _extract_pdf_text_from_bytes, fetch_tender_detail, fetch_tender_list
from state import filter_new, load_seen, save_seen

BASE_DIR = Path(__file__).parent
load_dotenv(BASE_DIR / ".env")

_SIGNER = URLSafeTimedSerializer(os.environ.get("SESSION_SECRET", "et-tools-2024-change-me"))
SETTINGS_FILE = BASE_DIR / "data" / "settings.json"
KNOWLEDGE_FILE = BASE_DIR / "data" / "shared" / "knowledge.json"
_excel_cache: dict[str, bytes] = {}

app = FastAPI(title="Electra Target Tools")
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

@app.on_event("startup")
async def _init_dirs():
    for d in ["data/shared", "data/users"]:
        (BASE_DIR / d).mkdir(parents=True, exist_ok=True)


# ── Session helpers ───────────────────────────────────────────────────────────

def _sign(username: str) -> str:
    return _SIGNER.dumps(username)

def _unsign(token: Optional[str]) -> Optional[str]:
    if not token:
        return None
    try:
        return _SIGNER.loads(token, max_age=30 * 86400)
    except (BadSignature, SignatureExpired, Exception):
        return None

def auth(et_session: Optional[str] = Cookie(default=None)) -> str:
    u = _unsign(et_session)
    if not u:
        raise HTTPException(401, "לא מחובר")
    return u

def _load_users() -> dict:
    with open(BASE_DIR / "users.yaml", encoding="utf-8") as f:
        return yaml.load(f, SafeLoader)


# ── Data helpers ──────────────────────────────────────────────────────────────

def _udir(u: str) -> Path:
    p = BASE_DIR / "data" / "users" / u
    p.mkdir(parents=True, exist_ok=True)
    return p

def _rj(path: Path, default):
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            pass
    return default

def _wj(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def _load_settings() -> dict:
    d = {
        "scraper": {
            "base_url": "https://mr.gov.il/ilgstorefront/he/search/?q=:relevance&inContract=false",
            "page_load_timeout_ms": 30000, "max_tenders_per_run": 50, "days_back": 7,
        },
        "budget": {"min_annual_ils": 500000, "max_annual_ils": 50000000},
        "industries": ["ניקיון", "אחזקה", "שמירה ואבטחה", "כוח אדם", "שירותי עזר", "קייטרינג"],
        "labor_costs": {"simple_monthly_ils": 6500, "simple_hourly_ils": 35, "social_expense_multiplier": 1.25},
    }
    if not SETTINGS_FILE.exists():
        _wj(SETTINGS_FILE, d)
        return d
    loaded = _rj(SETTINGS_FILE, d)
    for k, v in d.items():
        loaded.setdefault(k, v)
    return loaded

def _load_knowledge() -> list:
    return _rj(KNOWLEDGE_FILE, [])

def _save_knowledge(g: list):
    KNOWLEDGE_FILE.parent.mkdir(parents=True, exist_ok=True)
    _wj(KNOWLEDGE_FILE, g)

def _append_history(result: dict, u: str):
    h = _rj(_udir(u) / "history.json", [])
    if result["tender_id"] not in {r["tender_id"] for r in h}:
        h.append(result)
        _wj(_udir(u) / "history.json", h)

def _client() -> _anthropic.Anthropic:
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        raise HTTPException(500, "ANTHROPIC_API_KEY לא מוגדר")
    return _anthropic.Anthropic(api_key=key)


# ── Auth routes ───────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def index():
    return HTMLResponse(
        (BASE_DIR / "static" / "index.html").read_text(encoding="utf-8"),
        headers={"Cache-Control": "no-cache, no-store, must-revalidate"},
    )

@app.post("/api/login")
async def login(resp: Response, username: str = Form(...), password: str = Form(...)):
    data = _load_users()
    user = data["credentials"]["usernames"].get(username)
    if not user or not bcrypt.checkpw(password.encode(), user["password"].encode()):
        raise HTTPException(401, "שם משתמש או סיסמה שגויים")
    resp.set_cookie("et_session", _sign(username), max_age=30 * 86400, httponly=True, samesite="lax")
    return {"username": username, "name": user.get("name", username),
            "role": user.get("role", "user"), "apps": user.get("apps", ["tender_scanner", "shift_comparison"])}

@app.post("/api/logout")
async def logout(resp: Response):
    resp.delete_cookie("et_session")
    return {"ok": True}

@app.get("/api/me")
async def me(u: str = Depends(auth)):
    users = _load_users()["credentials"]["usernames"]
    info = users.get(u, {})
    return {"username": u, "name": info.get("name", u),
            "role": info.get("role", "user"), "apps": info.get("apps", ["tender_scanner", "shift_comparison"])}


# ── Settings & knowledge ──────────────────────────────────────────────────────

@app.get("/api/settings")
async def get_settings(_: str = Depends(auth)):
    return _load_settings()

@app.post("/api/settings")
async def post_settings(body: dict, _: str = Depends(auth)):
    _wj(SETTINGS_FILE, body)
    return {"ok": True}

@app.get("/api/knowledge")
async def get_knowledge(_: str = Depends(auth)):
    return _load_knowledge()

@app.post("/api/knowledge")
async def post_knowledge(body: dict, _: str = Depends(auth)):
    _save_knowledge(body.get("guidelines", []))
    return {"ok": True}

@app.delete("/api/knowledge/{idx}")
async def del_knowledge(idx: int, _: str = Depends(auth)):
    k = _load_knowledge()
    if 0 <= idx < len(k):
        k.pop(idx)
        _save_knowledge(k)
    return k


# ── Tender data ───────────────────────────────────────────────────────────────

@app.get("/api/history")
async def get_history(u: str = Depends(auth)):
    return _rj(_udir(u) / "history.json", [])

@app.patch("/api/history/{tid}")
async def patch_history(tid: str, body: dict, u: str = Depends(auth)):
    path = _udir(u) / "history.json"
    h = _rj(path, [])
    for entry in h:
        if entry.get("tender_id") == tid:
            entry.update(body)
            break
    _wj(path, h)
    return {"ok": True}

@app.delete("/api/history/{tid}")
async def delete_history(tid: str, u: str = Depends(auth)):
    path = _udir(u) / "history.json"
    h = _rj(path, [])
    h = [r for r in h if r.get("tender_id") != tid]
    _wj(path, h)
    # also remove from seen so it can be re-scanned
    seen_path = _udir(u) / "seen.json"
    seen = load_seen(seen_path)
    seen.discard(tid)
    save_seen(seen, seen_path)
    return {"ok": True}

@app.get("/api/last-scan")
async def get_last_scan(u: str = Depends(auth)):
    return _rj(_udir(u) / "last_scan.json", [])

@app.get("/api/favorites")
async def get_favorites(u: str = Depends(auth)):
    return _rj(_udir(u) / "favorites.json", [])

@app.post("/api/favorites/{tid}")
async def toggle_fav(tid: str, u: str = Depends(auth)):
    favs: list = _rj(_udir(u) / "favorites.json", [])
    if tid in favs:
        favs.remove(tid)
    else:
        favs.append(tid)
    _wj(_udir(u) / "favorites.json", favs)
    return favs


# ── Scan (SSE) ────────────────────────────────────────────────────────────────

@app.get("/api/scan-test")
async def scan_test(u: str = Depends(auth)):
    import traceback
    try:
        settings = _load_settings()
        from playwright.async_api import async_playwright
        async with async_playwright() as pw:
            browser = await pw.chromium.launch(headless=True)
            context = await browser.new_context(locale="he-IL", extra_http_headers={"Accept-Language": "he-IL,he;q=0.9"})
            page = await context.new_page()
            await page.goto(settings["scraper"]["base_url"], timeout=30000, wait_until="domcontentloaded")
            await page.wait_for_timeout(3000)
            final_url = page.url
            items = await page.query_selector_all("div.result-container")
            # check show-more button with various selectors
            show_more = None
            show_more_sel = None
            for sel in ["button.show-more-button","button:has-text('הצג עוד')","a:has-text('הצג עוד')",".show-more","[class*='show-more']"]:
                el = await page.query_selector(sel)
                if el:
                    show_more = await el.inner_text()
                    show_more_sel = sel
                    break
            # get first 3 titles
            titles = []
            for item in items[:3]:
                el = await item.query_selector("h2")
                if el: titles.append(await el.inner_text())
            await browser.close()
            return {"final_url": final_url, "items_found": len(items), "show_more_btn": show_more, "show_more_sel": show_more_sel, "sample_titles": titles}
    except Exception as e:
        return {"ok": False, "error": str(e), "trace": traceback.format_exc()[-1000:]}

@app.get("/api/scan")
async def scan(u: str = Depends(auth), skip_seen: bool = Query(True)):
    settings = _load_settings()

    async def gen():
        import traceback
        try:
            client = _client()
            loop = asyncio.get_running_loop()
        except Exception as e:
            yield f"data: {json.dumps({'type':'error','msg':f'שגיאת אתחול: {traceback.format_exc()[-400:]}'})}\n\n"
            return

        yield f"data: {json.dumps({'type':'status','msg':'מתחבר ל-mr.gov.il...'})}\n\n"
        try:
            import traceback
            tender_list = await asyncio.wait_for(fetch_tender_list(settings), timeout=120)
        except asyncio.TimeoutError:
            yield f"data: {json.dumps({'type':'error','msg':'timeout — הסריקה לקחה יותר מ-120 שניות'})}\n\n"
            return
        except Exception as e:
            yield f"data: {json.dumps({'type':'error','msg': traceback.format_exc()[-500:]})}\n\n"
            return

        seen_path = _udir(u) / "seen.json"
        seen = load_seen(seen_path) if skip_seen else set()
        new = filter_new(tender_list, seen) if skip_seen else tender_list

        yield f"data: {json.dumps({'type':'status','msg':f'נמצאו {len(tender_list)} מכרזים, {len(new)} חדשים לניתוח'})}\n\n"
        yield f"data: {json.dumps({'type':'count','total':len(tender_list),'new':len(new)})}\n\n"

        if not new:
            yield f"data: {json.dumps({'type':'complete','count':0,'total':len(tender_list)})}\n\n"
            return

        results = []
        for i, meta in enumerate(new):
            yield f"data: {json.dumps({'type':'progress','i':i+1,'total':len(new),'title':meta['title'][:80]})}\n\n"
            yield ": keepalive\n\n"

            try:
                tender = await asyncio.wait_for(fetch_tender_detail(meta, settings), timeout=90)
                tender.raw_metadata.update({"publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date","")})

                if not tender.pdf_text:
                    result = {
                        "tender_id": tender.tender_id, "title": tender.title, "url": tender.url,
                        "publisher": tender.publisher, "deadline": tender.deadline,
                        "publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date",""),
                        "has_pdf": False, "analysis": "NO_PDF",
                    }
                else:
                    knowledge = _load_knowledge()
                    result = await loop.run_in_executor(
                        None, lambda t=tender: analyze_tender(t, settings, client, knowledge=knowledge)
                    )
            except asyncio.TimeoutError:
                result = {
                    "tender_id": meta["tender_id"], "title": meta["title"], "url": meta["url"],
                    "publisher": "", "deadline": "",
                    "publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date",""),
                    "has_pdf": False, "analysis": "שגיאה: timeout",
                }
            except Exception as e:
                result = {
                    "tender_id": meta["tender_id"], "title": meta["title"], "url": meta["url"],
                    "publisher": "", "deadline": "",
                    "publish_date": meta.get("publish_date",""), "update_date": meta.get("update_date",""),
                    "has_pdf": False, "analysis": f"שגיאה: {e}",
                }

            results.append(result)
            _append_history(result, u)
            seen.add(meta["tender_id"])
            if skip_seen:
                save_seen(seen, seen_path)
            _wj(_udir(u) / "last_scan.json", results)

            yield f"data: {json.dumps({'type':'result','data':result,'pct':(i+1)/len(new)})}\n\n"

        yield f"data: {json.dumps({'type':'complete','count':len(results),'total':len(tender_list)})}\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream",
                              headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── Analyze single URL ────────────────────────────────────────────────────────

@app.post("/api/analyze")
async def analyze_url(body: dict, u: str = Depends(auth)):
    url = body.get("url", "").strip()
    if not url:
        raise HTTPException(400, "URL נדרש")
    settings = _load_settings()
    tid = _extract_id_from_url(url)
    meta = {"tender_id": tid, "title": url, "url": url, "publish_date": "", "update_date": ""}
    try:
        tender = await asyncio.wait_for(fetch_tender_detail(meta, settings), timeout=90)
        if not tender.title or tender.title.startswith("http"):
            tender.title = f"מכרז {tid}"
        tender.raw_metadata.update({"publish_date": "", "update_date": ""})
        if not tender.pdf_text:
            result = {
                "tender_id": tid, "title": tender.title, "url": url,
                "publisher": tender.publisher, "deadline": tender.deadline,
                "publish_date": "", "update_date": "", "has_pdf": False, "analysis": "NO_PDF",
            }
        else:
            client = _client()
            loop = asyncio.get_running_loop()
            knowledge = _load_knowledge()
            result = await loop.run_in_executor(
                None, lambda t=tender: analyze_tender(t, settings, client, knowledge=knowledge)
            )
        _append_history(result, u)
        return result
    except asyncio.TimeoutError:
        raise HTTPException(408, "timeout")
    except Exception as e:
        raise HTTPException(500, str(e))


# ── Analyze uploaded PDF ──────────────────────────────────────────────────────

@app.post("/api/analyze-pdf")
async def analyze_pdf_upload(pdf: UploadFile = File(...), u: str = Depends(auth)):
    body = await pdf.read()
    if not body:
        raise HTTPException(400, "קובץ ריק")
    pdf_text = _extract_pdf_text_from_bytes(body)
    if not pdf_text.strip():
        raise HTTPException(422, "לא ניתן לחלץ טקסט מהקובץ")
    settings = _load_settings()
    client = _client()
    knowledge = _load_knowledge()
    filename = pdf.filename or "מכרז"
    title = filename.replace(".pdf", "").replace("_", " ")
    tender = Tender(
        tender_id=f"pdf_{int(time.time())}",
        title=title,
        url="",
        publisher="",
        deadline="",
        raw_metadata={"publish_date": "", "update_date": ""},
        pdf_text=pdf_text,
    )
    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(
        None, lambda: analyze_tender(tender, settings, client, knowledge=knowledge)
    )
    _append_history(result, u)
    return result


# ── Chat ──────────────────────────────────────────────────────────────────────

@app.post("/api/chat")
async def chat(body: dict, _: str = Depends(auth)):
    tender = body.get("tender", {})
    history = body.get("history", [])
    client = _client()
    system = f"""אתה יועץ עסקי שניתח מכרז עבור Electra Target.
כותרת: {tender.get('title','')}
מפרסם: {tender.get('publisher','לא ידוע')}
מועד הגשה: {tender.get('deadline','לא צוין')}
ניתוח: {tender.get('analysis','')}
ענה בעברית בלבד. היה ממוקד ותמציתי."""
    msgs = [{"role": m["role"], "content": m["content"]} for m in history]
    loop = asyncio.get_running_loop()
    try:
        resp = await loop.run_in_executor(
            None,
            lambda: client.messages.create(
                model="claude-sonnet-4-6", max_tokens=1024, system=system, messages=msgs
            )
        )
        return {"reply": resp.content[0].text}
    except Exception as e:
        raise HTTPException(500, str(e))


# ── Re-analyze with clarification answers ─────────────────────────────────────

@app.post("/api/tender/reanalyze-with-answers")
async def reanalyze_with_answers(
    answers: UploadFile = File(...),
    tender_data: str = Form(...),
    u: str = Depends(auth)
):
    import json as _json
    td = _json.loads(tender_data)
    answers_text = (await answers.read()).decode("utf-8", errors="ignore")
    settings = _load_settings()
    client = _client()
    knowledge = _load_knowledge()
    tender = Tender(
        tender_id=td.get("tender_id",""),
        title=td.get("title",""),
        url=td.get("url",""),
        publisher=td.get("publisher",""),
        deadline=td.get("deadline",""),
        raw_metadata={"publish_date": td.get("publish_date",""), "update_date": td.get("update_date","")},
        pdf_text=td.get("pdf_text",""),
    )
    feedback = [f"תשובות הבהרה שהתקבלו:\n{answers_text}"]
    loop = asyncio.get_running_loop()
    result = await loop.run_in_executor(
        None, lambda: analyze_tender(tender, settings, client, knowledge=knowledge, session_feedback=feedback)
    )
    _append_history(result, u)
    return result


# ── Generate clarification questions ──────────────────────────────────────────

@app.post("/api/tender/generate-questions")
async def generate_questions(body: dict, _: str = Depends(auth)):
    r = body
    client = _client()
    knowledge = _load_knowledge()
    prompt = f"""אתה יועץ עסקי של Electra Target. קראת את הניתוח הבא של מכרז ממשלתי.

כותרת: {r.get('title','')}
מפרסם: {r.get('publisher','')}
מועד הגשה: {r.get('deadline','')}
ניתוח: {r.get('analysis','')}

צור את כל שאלות ההבהרה שיש לשלוח למפרסם המכרז — כמה שצריך, עד 50 שאלות לכל היותר.
כלול כל שאלה חשובה שעולה מהמסמך. אל תגביל את עצמך למספר קבוע.

פלט כל שאלה בפורמט הבא (עמודה מופרדת ב-|):
מספר|עמוד|סעיף|שאלה

- עמוד: מספר עמוד רלוונטי במסמך אם ידוע, אחרת: כללי
- סעיף: מספר סעיף רלוונטי אם ידוע, אחרת: כללי
- שאלה: טקסט השאלה בעברית

דוגמה:
1|כללי|כללי|האם נדרש רישיון עסק?
2|5|3.2|מה תקופת האחריות על הציוד?

כתוב את כל השאלות בפורמט זה בלבד, ללא כותרות נוספות."""

    system = SYSTEM_PROMPT
    if knowledge:
        system += "\n\nתובנות שנצברו:\n" + "\n".join(f"- {g}" for g in knowledge)

    loop = asyncio.get_running_loop()
    try:
        resp = await loop.run_in_executor(
            None,
            lambda: client.messages.create(
                model="claude-sonnet-4-6", max_tokens=4096, system=system,
                messages=[{"role": "user", "content": prompt}]
            )
        )
        return {"questions": resp.content[0].text}
    except Exception as e:
        raise HTTPException(500, str(e))


# ── Export Excel (financial analysis) ─────────────────────────────────────────

@app.post("/api/tender/export-excel")
async def export_excel(body: dict, _: str = Depends(auth)):
    import io, re
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    r = body
    analysis = r.get("analysis", "")

    # Split analysis into sections
    fin_keys = ["הערכת היקף כספי", "ערך שנתי", "בסיס לחישוב", "כוח אדם נדרש", "אורך חוזה", "המלצה", "אתגרים", "סיכונים"]
    sections: dict[str, list[str]] = {}
    current_key = None
    for line in analysis.splitlines():
        clean = line.strip().replace("**","").replace("#","").strip()
        if not clean: continue
        matched = next((k for k in fin_keys if k in clean), None)
        if matched:
            current_key = matched
            sections.setdefault(current_key, [])
        if current_key:
            sections[current_key].append(clean)

    wb = Workbook()
    ws = wb.active
    ws.title = "ניתוח פיננסי"
    ws.sheet_view.rightToLeft = True

    hfill = PatternFill("solid", fgColor="1E3A5F")
    sfill = PatternFill("solid", fgColor="2563EB")
    thfill= PatternFill("solid", fgColor="374151")
    afill = PatternFill("solid", fgColor="EEF3FA")
    wfill = PatternFill("solid", fgColor="FFFFFF")
    thin  = Border(left=Side(style='thin',color='CCCCCC'), right=Side(style='thin',color='CCCCCC'),
                   top=Side(style='thin',color='CCCCCC'),  bottom=Side(style='thin',color='CCCCCC'))

    def is_md_table(line):
        return line.startswith('|') and line.endswith('|')

    def is_separator(line):
        return re.fullmatch(r'[\|\-\s:]+', line) is not None

    def parse_md_row(line):
        return [c.strip() for c in line.strip('|').split('|')]

    def rtl_align(horizontal="right", center=False):
        return Alignment(horizontal="center" if center else horizontal,
                         vertical="center", wrap_text=True)

    def style_cell(c, bold=False, fill=None, center=False):
        c.font = Font(bold=bold, name="Arial", size=10)
        if fill: c.fill = fill
        c.border = thin
        c.alignment = rtl_align(center=center)

    def hrow(label, row, fill=None, ncols=2):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=ncols)
        c = ws.cell(row, 1, label)
        c.font = Font(bold=True, color="FFFFFF", size=12, name="Arial")
        c.fill = fill or hfill
        c.alignment = rtl_align()
        ws.row_dimensions[row].height = 24

    def drow(label, value, row, alt=False, ncols=2):
        fill = afill if alt else wfill
        c1 = ws.cell(row, 1, label or "")
        style_cell(c1, bold=bool(label), fill=fill)
        c2 = ws.cell(row, 2, str(value) if value else "")
        style_cell(c2, fill=fill)
        ws.row_dimensions[row].height = max(18, min(90, len(str(value or ""))//3+15))

    def write_section_lines(lines, start_row, max_cols):
        r = start_row
        alt = False
        i = 0
        while i < len(lines):
            line = lines[i]
            if is_md_table(line):
                # collect table block
                tbl_lines = []
                while i < len(lines) and (is_md_table(lines[i]) or is_separator(lines[i])):
                    tbl_lines.append(lines[i]); i += 1
                data_rows = [parse_md_row(l) for l in tbl_lines if not is_separator(l)]
                if not data_rows: continue
                ncols = max(len(row) for row in data_rows)
                # set/extend column widths
                for ci in range(1, ncols+1):
                    col_letter = ws.cell(r, ci).column_letter
                    ws.column_dimensions[col_letter].width = max(
                        ws.column_dimensions[col_letter].width, 18)
                for ri, dr in enumerate(data_rows):
                    is_hdr = ri == 0
                    row_fill = thfill if is_hdr else (afill if ri%2==0 else wfill)
                    for ci, val in enumerate(dr):
                        c = ws.cell(r, ci+1, val)
                        style_cell(c, bold=is_hdr, fill=row_fill,
                                   center=(ci < len(dr)-1))
                        if is_hdr: c.font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
                    ws.row_dimensions[r].height = 20
                    r += 1
            else:
                clean = line.replace("**","").replace("#","").strip()
                if clean and not is_separator(clean):
                    c = ws.cell(r, 1, clean)
                    style_cell(c, fill=(afill if alt else wfill))
                    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=max_cols)
                    ws.row_dimensions[r].height = max(18, min(60, len(clean)//5+15))
                    alt = not alt
                    r += 1
                i += 1
        return r

    # Auto-detect column count from tables in analysis
    max_cols = 2
    for line in analysis.splitlines():
        if is_md_table(line.strip()):
            max_cols = max(max_cols, line.count('|') - 1)
    max_cols = min(max_cols, 8)
    for ci in range(1, max_cols+1):
        letter = ws.cell(1, ci).column_letter
        ws.column_dimensions[letter].width = 20
    ws.column_dimensions[ws.cell(1,1).column_letter].width = 30

    row = 1
    hrow(f"ניתוח פיננסי: {r.get('title','')}", row, ncols=max_cols); row += 1
    for label, val, alt in [
        ("מפרסם", r.get("publisher",""), False),
        ("מועד הגשה", r.get("deadline",""), True),
        ("תאריך פרסום", r.get("publish_date",""), False),
    ]:
        drow(label, val, row, alt, ncols=max_cols); row += 1
    row += 1

    finance_order = ["הערכת היקף כספי", "ערך שנתי", "בסיס לחישוב", "כוח אדם נדרש", "אורך חוזה", "המלצה", "אתגרים", "סיכונים"]
    hrow("ניתוח פיננסי מפורט", row, sfill, ncols=max_cols); row += 1
    shown = set()
    for key in finance_order:
        lines = sections.get(key)
        if not lines or key in shown: continue
        shown.add(key)
        # Section sub-header
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=max_cols)
        c = ws.cell(row, 1, key)
        c.font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
        c.fill = PatternFill("solid", fgColor="2B5C9E")
        c.alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row].height = 20; row += 1
        row = write_section_lines(lines[1:], row, max_cols)  # skip key header line

    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    tid = r.get("tender_id", "tender")
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''%D7%A4%D7%99%D7%A0%D7%A0%D7%A1%D7%99_{tid}.xlsx"})


# ── Export Word (full analysis) ────────────────────────────────────────────────

@app.post("/api/tender/export-word")
async def export_word(body: dict, _: str = Depends(auth)):
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r = body
    analysis  = r.get("analysis", "")
    questions = r.get("questions", "")
    knowledge = _load_knowledge()

    doc = Document()
    # Set document-level RTL
    sectPr = doc.sections[0]._sectPr
    sectPr.append(OxmlElement('w:bidi'))

    def set_rtl_para(par):
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = par._p.get_or_add_pPr()
        b = OxmlElement('w:bidi'); b.set(qn('w:val'), '1'); pPr.append(b)
        jc = OxmlElement('w:jc');  jc.set(qn('w:val'), 'right'); pPr.append(jc)

    def set_rtl_run(run):
        run.font.cs_name = "Arial"
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)
        rPr.append(OxmlElement('w:cs'))

    def h(text, level=1):
        par = doc.add_heading(text, level=level)
        set_rtl_para(par)
        for run in par.runs:
            run.font.name = "Arial"
            set_rtl_run(run)
            if level == 1:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        return par

    from docx.shared import Cm
    from docx.oxml import OxmlElement as OE

    def add_run_inline(par, text):
        """Add run with inline **bold** parsing."""
        import re
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            if not part: continue
            run = par.add_run(part)
            run.font.name = "Arial"; run.font.size = Pt(11)
            run.bold = (i % 2 == 1)
            set_rtl_run(run)

    def p(text, bold=False, bullet=False, size=11):
        par = doc.add_paragraph()
        set_rtl_para(par)
        if bullet:
            pPr = par._p.get_or_add_pPr()
            ind = OE('w:ind'); ind.set(qn('w:right'), '360'); pPr.append(ind)
        if bold:
            run = par.add_run(str(text))
            run.font.name = "Arial"; run.font.cs_name = "Arial"
            run.font.size = Pt(size); run.bold = True
            set_rtl_run(run)
        else:
            add_run_inline(par, str(text))
        return par

    def render_md(text):
        """Render markdown text into the Word document."""
        import re
        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            # Heading
            hm = re.match(r'^(#{1,3})\s+(.*)', stripped)
            if hm:
                level = min(len(hm.group(1)) + 1, 3)
                h(hm.group(2).replace('**',''), level)
                i += 1; continue
            # Table
            if stripped.startswith('|') and stripped.endswith('|'):
                tbl_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    tbl_lines.append(lines[i].strip()); i += 1
                data_rows = [r for r in tbl_lines if not re.fullmatch(r'[\|\-\s:]+', r)]
                if not data_rows: continue
                parsed = [[c.strip() for c in r.strip('|').split('|')] for r in data_rows]
                ncols = max(len(r) for r in parsed)
                tbl = doc.add_table(rows=len(parsed), cols=ncols)
                tbl.style = 'Table Grid'
                for ri, row in enumerate(parsed):
                    for ci, val in enumerate(row):
                        cell = tbl.rows[ri].cells[ci]
                        cell.paragraphs[0].clear()
                        cp = cell.paragraphs[0]
                        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        pPr = cp._p.get_or_add_pPr()
                        b = OE('w:bidi'); b.set(qn('w:val'),'1'); pPr.append(b)
                        run = cp.add_run(val)
                        run.font.name = "Arial"; run.font.size = Pt(10)
                        run.bold = (ri == 0)
                        if ri == 0: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                        set_rtl_run(run)
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OE('w:shd')
                        shd.set(qn('w:fill'), '1E3A5F' if ri==0 else ('EEF3FA' if ri%2==0 else 'FFFFFF'))
                        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear')
                        tcPr.append(shd)
                continue
            # Bullet
            bm = re.match(r'^[-*+]\s+(.*)', stripped)
            if bm:
                p(f"• {bm.group(1)}", bullet=True); i += 1; continue
            # Numbered list
            nm = re.match(r'^\d+\.\s+(.*)', stripped)
            if nm:
                p(f"• {nm.group(1)}", bullet=True); i += 1; continue
            # Horizontal rule or separator — skip
            if re.fullmatch(r'[-_*]{2,}', stripped):
                i += 1; continue
            # Empty
            if not stripped:
                i += 1; continue
            # Normal
            p(stripped)
            i += 1

    h(f"ניתוח מכרז: {r.get('title','')}", 1)
    if r.get('publisher'):    p(f"מפרסם: {r['publisher']}")
    if r.get('publish_date'): p(f"תאריך פרסום: {r['publish_date']}")
    if r.get('deadline'):     p(f"מועד הגשה: {r['deadline']}")
    if r.get('update_date'):  p(f"תאריך עדכון: {r['update_date']}")
    doc.add_paragraph()

    render_md(analysis.replace('<','').replace('>',''))

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    tid = r.get("tender_id", "tender")
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''%D7%A0%D7%99%D7%AA%D7%95%D7%97_{tid}.docx"})


# ── Export questions to Word (table) ──────────────────────────────────────────

@app.post("/api/tender/export-questions-word")
async def export_questions_word(body: dict, _: str = Depends(auth)):
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r = body
    questions_raw = r.get("questions", "")

    doc = Document()
    sectPr = doc.sections[0]._sectPr
    sectPr.append(OxmlElement('w:bidi'))

    def set_rtl(par):
        pPr = par._p.get_or_add_pPr()
        b = OxmlElement('w:bidi'); b.set(qn('w:val'), '1'); pPr.append(b)
        jc = OxmlElement('w:jc');  jc.set(qn('w:val'), 'right'); pPr.append(jc)

    def set_rtl_run(run):
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)
        rPr.append(OxmlElement('w:cs'))
        run.font.cs_name = "Arial"

    title_p = doc.add_heading(f"שאלות הבהרה: {r.get('title','')}", 1)
    set_rtl(title_p)
    for run in title_p.runs:
        run.font.name = "Arial"; run.font.cs_name = "Arial"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        set_rtl_run(run)

    def meta_line(text):
        p = doc.add_paragraph(); set_rtl(p)
        run = p.add_run(text)
        run.font.name = "Arial"; run.font.cs_name = "Arial"; run.font.size = Pt(11)
        set_rtl_run(run)

    if r.get('publisher'):    meta_line(f"מפרסם: {r['publisher']}")
    if r.get('publish_date'): meta_line(f"תאריך פרסום: {r['publish_date']}")
    if r.get('deadline'):     meta_line(f"מועד הגשה: {r['deadline']}")
    if r.get('update_date'):  meta_line(f"תאריך עדכון: {r['update_date']}")
    doc.add_paragraph()

    # Parse pipe-separated questions
    rows = []
    for line in questions_raw.strip().splitlines():
        line = line.strip()
        if not line: continue
        parts = line.split("|")
        if len(parts) >= 4:
            rows.append((parts[0].strip(), parts[1].strip(), parts[2].strip(), "|".join(parts[3:]).strip()))
        elif line:
            num = str(len(rows)+1)
            rows.append((num, "כללי", "כללי", line.lstrip("0123456789. ")))

    # Column order as written in file — RTL doc renders col1 on the right
    # so מספר(col1) | עמוד(col2) | סעיף(col3) | שאלה(col4) displays correctly R→L
    headers_rtl  = ["שאלה", "סעיף", "עמוד", "מספר"]
    col_widths_rtl = [Cm(11), Cm(2.5), Cm(2), Cm(1.5)]

    def add_shd(tc, color):
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def cell_rtl_run(cell, text, bold=False, size=10, color=None):
        p = cell.paragraphs[0]; p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = p._p.get_or_add_pPr()
        b = OxmlElement('w:bidi'); b.set(qn('w:val'),'1'); pPr.append(b)
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'),'right'); pPr.append(jc)
        run = p.add_run(text)
        run.font.name = "Arial"; run.font.cs_name = "Arial"; run.font.size = Pt(size); run.bold = bold
        if color: run.font.color.rgb = color
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'),'1'); rPr.append(rtl)
        rPr.append(OxmlElement('w:cs'))

    table = doc.add_table(rows=1+len(rows), cols=4)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, (h_text, w) in enumerate(zip(headers_rtl, col_widths_rtl)):
        cell = hdr.cells[i]; cell.width = w
        add_shd(cell._tc, '1E3A5F')
        cell_rtl_run(cell, h_text, bold=True, size=11, color=RGBColor(0xFF,0xFF,0xFF))

    # Data rows
    for ri, (num, page, section, question) in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        fill_color = 'EEF3FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, text in enumerate([question, section, page, num]):
            cell = row_cells[ci]
            add_shd(cell._tc, fill_color)
            cell_rtl_run(cell, text, size=10)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    tid = r.get("tender_id", "tender")
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''%D7%A9%D7%90%D7%9C%D7%95%D7%AA_{tid}.docx"})


# ── Email digest ──────────────────────────────────────────────────────────────

@app.post("/api/send-email")
async def send_email_digest(body: dict, u: str = Depends(auth)):
    from emailer import send_digest, build_html_digest
    from datetime import date
    import re

    tenders   = body.get("tenders", [])
    min_level = body.get("min_level", "medium")
    recipient = body.get("to", "").strip() or _load_settings()["email"]["recipient"]

    level_rank = {"high": 0, "medium": 1, "low": 2}
    threshold  = level_rank.get(min_level, 0)

    def get_rank(analysis):
        first_line = (analysis or "").split("\n")[0]
        if "גבוהה" in first_line: return 0
        if "בינונית" in first_line: return 1
        return 2

    def extract_summary(analysis: str) -> str:
        """Extract the first סיכום section from analysis."""
        lines = analysis.splitlines()
        in_section = False
        result = []
        for line in lines:
            if re.search(r'סיכום', line):
                in_section = True
                continue
            if in_section:
                if re.match(r'^#{1,3}\s', line) and result:
                    break
                if line.strip():
                    result.append(line.strip().replace('**','').replace('#',''))
        return ' '.join(result[:5]) if result else analysis[:300]

    filtered = [t for t in tenders if get_rank(t.get("analysis","")) <= threshold]
    if not filtered:
        return {"ok": False, "msg": "לא נמצאו מכרזים ברמה שנבחרה"}

    app_url = "https://tender-scanner.up.railway.app"
    def badge(a):
        if "גבוהה" in a: return "🟢 גבוהה"
        if "בינונית" in a: return "🟡 בינונית"
        return "🔴 נמוכה"
    def badge_color(a):
        if "גבוהה" in a: return "#1a7a1a"
        if "בינונית" in a: return "#b36b00"
        return "#8b0000"

    cards = []
    for t in filtered:
        analysis = t.get("analysis","")
        summary  = extract_summary(analysis)
        color    = badge_color(analysis)
        cards.append(f"""
        <div style="border:1px solid #ddd;border-radius:8px;padding:16px;margin-bottom:20px;font-family:Arial,sans-serif;direction:rtl;text-align:right;">
          <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px">
            <h2 style="margin:0;font-size:16px;color:{color}">{t.get('title','')}</h2>
            <span style="background:{color};color:#fff;padding:3px 10px;border-radius:12px;font-size:12px;white-space:nowrap">{badge(analysis)}</span>
          </div>
          <p style="margin:4px 0;color:#555;font-size:13px">
            {t.get('publisher','') or ''}
            {' | פרסום: '+t['publish_date'] if t.get('publish_date') else ''}
            {' | הגשה: '+t['deadline'] if t.get('deadline') else ''}
          </p>
          <hr style="border:none;border-top:1px solid #eee;margin:10px 0">
          <p style="font-size:14px;line-height:1.7;margin:0 0 12px">{summary}</p>
          <div style="display:flex;gap:12px">
            <a href="{t.get('url','')}" style="color:#2563EB;font-size:13px">🔗 עמוד המכרז</a>
            <a href="{app_url}" style="color:#2563EB;font-size:13px">📊 ניתוח מלא במערכת</a>
          </div>
        </div>""")

    run_date = date.today().strftime("%d/%m/%Y")
    level_label = {"high":"גבוהה בלבד","medium":"בינונית וגבוהה","low":"כל הרמות"}.get(min_level,"")
    html = f"""<html><body style="background:#f5f5f5;padding:20px">
      <h1 style="font-family:Arial,sans-serif;direction:rtl;text-align:right;color:#1a1a2e">
        סריקת מכרזים — {run_date}
      </h1>
      <p style="font-family:Arial,sans-serif;direction:rtl;text-align:right;color:#555">
        {len(filtered)} מכרזים ברמה: {level_label}
      </p>
      {''.join(cards)}
      <p style="font-family:Arial,sans-serif;font-size:12px;color:#999;text-align:center;margin-top:30px">Electra Target Tools</p>
    </body></html>"""

    resend_key = os.environ.get("RESEND_API_KEY","")
    if not resend_key:
        return {"ok": False, "msg": "RESEND_API_KEY לא מוגדר בסביבה"}

    import urllib.request, json as _json
    subject = f"[Electra Target] {len(filtered)} מכרזים — {run_date}"
    try:
        payload = _json.dumps({"from": "Electra Target <onboarding@resend.dev>", "to": [recipient], "subject": subject, "html": html}).encode()
        req = urllib.request.Request("https://api.resend.com/emails", data=payload,
            headers={"Authorization": f"Bearer {resend_key}", "Content-Type": "application/json"}, method="POST")
        with urllib.request.urlopen(req, timeout=15) as resp:
            return {"ok": True, "count": len(filtered)}
    except urllib.error.HTTPError as e:
        body = e.read().decode()
        return {"ok": False, "msg": f"HTTP {e.code}: {body}"}
    except Exception as e:
        import traceback
        return {"ok": False, "msg": traceback.format_exc()[-600:]}


# ── Learning mode ─────────────────────────────────────────────────────────────

@app.post("/api/learn/analyze")
async def learn_analyze(body: dict, u: str = Depends(auth)):
    url = body.get("url", "").strip()
    settings = _load_settings()
    tid = _extract_id_from_url(url)
    meta = {"tender_id": tid, "title": url, "url": url, "publish_date": "", "update_date": ""}
    tender = await asyncio.wait_for(fetch_tender_detail(meta, settings), timeout=90)
    if not tender.title or tender.title.startswith("http"):
        tender.title = f"מכרז {tid}"
    td = {
        "tender_id": tender.tender_id, "title": tender.title, "url": tender.url,
        "publisher": tender.publisher, "deadline": tender.deadline,
        "pdf_text": tender.pdf_text, "has_pdf": bool(tender.pdf_text),
    }
    if not tender.pdf_text:
        result = {"title": tender.title, "url": url, "publisher": tender.publisher,
                  "deadline": tender.deadline, "has_pdf": False, "analysis": "NO_PDF"}
    else:
        client = _client()
        loop = asyncio.get_running_loop()
        knowledge = _load_knowledge()
        result = await loop.run_in_executor(
            None, lambda t=tender: analyze_tender(t, settings, client, knowledge=knowledge)
        )
    return {"tender": td, "result": result}

@app.post("/api/learn/reanalyze")
async def learn_reanalyze(body: dict, _: str = Depends(auth)):
    td = body.get("tender", {})
    feedback = body.get("feedback", [])
    settings = _load_settings()
    tender_obj = Tender(
        tender_id=td["tender_id"], title=td["title"], url=td["url"],
        publisher=td.get("publisher",""), deadline=td.get("deadline",""),
        pdf_text=td.get("pdf_text",""), raw_metadata={"publish_date":"","update_date":""},
    )
    client = _client()
    loop = asyncio.get_running_loop()
    knowledge = _load_knowledge()
    result = await loop.run_in_executor(
        None,
        lambda: analyze_tender(tender_obj, settings, client, knowledge=knowledge, session_feedback=feedback)
    )
    return result

@app.post("/api/learn/save")
async def learn_save(body: dict, _: str = Depends(auth)):
    title = body.get("title", "")
    feedback = body.get("feedback", [])
    client = _client()
    loop = asyncio.get_running_loop()
    new_g = await loop.run_in_executor(None, lambda: distill_knowledge(title, feedback, client))
    existing = _load_knowledge()
    existing.extend(new_g)
    _save_knowledge(existing)
    return {"new": new_g, "total": len(existing)}


# ── Shift comparison ──────────────────────────────────────────────────────────

_DEFAULT_SHIFTS_CFG = {
    "excel_columns": {"date": "B", "start_time": "C", "end_time": "D", "worker_name": "F"},
    "excel_has_header": True,
    "rules": {"gap_threshold_minutes": 30, "default_start_time": "10:00"},
    "aliases": {},
    "managers": [],
    "ignored_names": [],
}

def _shifts_cfg_path(u: str) -> Path:
    return _udir(u) / "shifts_config.json"

def _load_shifts_cfg(u: str) -> dict:
    return _rj(_shifts_cfg_path(u), dict(_DEFAULT_SHIFTS_CFG))

def _save_shifts_cfg(u: str, cfg: dict):
    _wj(_shifts_cfg_path(u), cfg)

@app.get("/api/shifts/config")
async def shifts_config(u: str = Depends(auth)):
    return _load_shifts_cfg(u)

@app.post("/api/shifts/config")
async def save_shifts_config(body: dict, u: str = Depends(auth)):
    _save_shifts_cfg(u, body)
    return {"ok": True}

@app.post("/api/shifts/validate")
async def shifts_validate(_: str = Depends(auth), message: str = Form(...)):
    return {"suspicious": find_suspicious_lines(message)}

GMAIL_EXCEL_COLUMNS = {
    "date":        "A",
    "worker_name": "D",
    "start_time":  "J",
    "end_time":    "P",
}

@app.post("/api/shifts/compare")
async def shifts_compare(
    u: str = Depends(auth),
    message: str = Form(...),
    excel: UploadFile = File(...),
    source: str = Form("upload"),
):
    cfg = _load_shifts_cfg(u)

    # Build aliases and managers from workers table (takes precedence over cfg)
    workers = _load_workers(u)
    worker_aliases = {
        w["nickname"]: w["full_name"]
        for w in workers
        if w.get("nickname") and w.get("full_name") and w["nickname"] != w["full_name"]
    }
    worker_managers = [
        w["full_name"] for w in workers if w.get("rank") == "manager" and w.get("full_name")
    ]
    known_worker_names = {w["full_name"] for w in workers if w.get("full_name")}

    compare_cfg = dict(cfg)
    merged_aliases = dict(cfg.get("aliases", {}))
    merged_aliases.update(worker_aliases)
    compare_cfg["aliases"] = merged_aliases
    if worker_managers:
        compare_cfg["managers"] = worker_managers
    compare_cfg["known_workers"] = known_worker_names
    # Same matcher the saved-shift enrichment and reports use, so the comparison
    # and everything downstream agree on who each line belongs to
    compare_cfg["worker_matcher"] = _worker_matcher(u)
    compare_cfg["branch_map"] = _build_branch_map(_load_branches(u))

    excel_bytes = await excel.read()
    if source == "gmail":
        parse_cfg = dict(cfg)
        parse_cfg["excel_columns"] = GMAIL_EXCEL_COLUMNS
        parse_cfg["excel_has_header"] = True
    else:
        parse_cfg = cfg

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(excel_bytes)
        tmp_path = Path(tmp.name)
    try:
        msg_entries = parse_message(message)
        excel_entries = parse_excel(str(tmp_path), parse_cfg)
        results = compare(msg_entries, excel_entries, compare_cfg)

        out_path = Path(tempfile.mktemp(suffix=".xlsx"))
        export_to_excel(results, str(out_path))
        excel_bytes = out_path.read_bytes()
        out_path.unlink(missing_ok=True)

        token = secrets.token_urlsafe(16)
        _excel_cache[token] = excel_bytes

        def _f(v):
            if hasattr(v, "strftime"):
                # time objects have .hour but not .year; date/datetime/Timestamp have both
                return v.strftime("%H:%M") if (hasattr(v, "hour") and not hasattr(v, "year")) else v.strftime("%d/%m/%Y")
            return str(v) if v else ""

        return {
            "results": [
                {
                    "status": r["status"],
                    "worker_name": r.get("worker_name", ""),
                    "worker_key": r.get("worker_key", ""),
                    "workplace": r.get("workplace", ""),
                    "branch_key": r.get("branch_key", ""),
                    "date": _f(r.get("date")),
                    "start_time": _f(r.get("start_time")),
                    "end_time": _f(r.get("end_time")),
                    "hours": round(r["hours"], 2) if r.get("hours") is not None else None,
                    "sales": str(r.get("sales", "")) if r.get("sales") else "",
                    "notes": r.get("notes", ""),
                }
                for r in results
            ],
            "token": token,
        }
    finally:
        tmp_path.unlink(missing_ok=True)

@app.get("/api/shifts/download/{token}")
async def shifts_download(token: str, _: str = Depends(auth)):
    data = _excel_cache.get(token)
    if not data:
        raise HTTPException(404, "קובץ לא נמצא")
    return Response(content=data,
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": "attachment; filename=comparison.xlsx"})


# ── Saved Shifts ─────────────────────────────────────────────────────────────

def _saved_shifts_path(u: str) -> Path:
    return _udir(u) / "saved_shifts.json"

def _load_saved_shifts(u: str) -> list:
    return _rj(_saved_shifts_path(u), [])

def _save_saved_shifts(u: str, rows: list):
    _saved_shifts_path(u).write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")

def _worker_matcher(u: str):
    """The shared NameMatcher for this user (workers table + nicknames + cfg aliases)."""
    return build_worker_matcher(_load_workers(u), _load_shifts_cfg(u).get("aliases", {}))

def _worker_resolver(u: str):
    """resolve(name) -> workers-table full name, or "" when unknown/ambiguous."""
    return _worker_matcher(u).resolve_or_blank

def _enrich_shift_row(row: dict, resolve_worker, branch_map: dict, prev: dict = None):
    """Recompute worker_key + branch_key from the row's text fields.
    Auto-resolution wins, except when the incoming value differs from the stored
    one (an explicit manual connection) — then the manual value is kept. When
    auto-resolution fails, an existing manual value survives."""
    prev = prev or {}
    auto_w   = resolve_worker(row.get("worker_name", ""))
    manual_w = (row.get("worker_key") or "").strip()
    if manual_w and manual_w != (prev.get("worker_key") or ""):
        row["worker_key"] = manual_w
    else:
        row["worker_key"] = auto_w or manual_w
    auto_b   = branch_map.get(normalize_match_text(row.get("workplace") or ""), "")
    manual_b = (row.get("branch_key") or "").strip()
    if manual_b and manual_b != (prev.get("branch_key") or ""):
        row["branch_key"] = manual_b
    else:
        row["branch_key"] = auto_b or manual_b

@app.post("/api/shifts/saved")
async def save_shifts(request: Request, u: str = Depends(auth)):
    import uuid as _uuid
    body = await request.json()
    existing = _load_saved_shifts(u)
    resolve = _worker_resolver(u)
    branch_map = _build_branch_map(_load_branches(u))
    for row in body:
        row["id"] = str(_uuid.uuid4())
        _enrich_shift_row(row, resolve, branch_map)
    existing.extend(body)
    _save_saved_shifts(u, existing)
    return {"ok": True, "saved": len(body)}

@app.post("/api/shifts/saved/recalculate")
async def recalculate_saved_shifts(u: str = Depends(auth)):
    rows = _load_saved_shifts(u)
    resolve = _worker_resolver(u)
    branch_map = _build_branch_map(_load_branches(u))
    for r in rows:
        _enrich_shift_row(r, resolve, branch_map, prev=r)
    _save_saved_shifts(u, rows)
    return {
        "ok": True,
        "total": len(rows),
        "workers_matched":  sum(1 for r in rows if r.get("worker_key")),
        "branches_matched": sum(1 for r in rows if r.get("branch_key")),
    }

@app.get("/api/shifts/saved")
async def get_saved_shifts(u: str = Depends(auth), date: str = Query(None), name: str = Query(None)):
    rows = _load_saved_shifts(u)
    if date:
        rows = [r for r in rows if date in str(r.get("date", ""))]
    if name:
        nl = name.lower()
        rows = [r for r in rows if nl in str(r.get("worker_name", "")).lower()]
    rows.sort(key=lambda r: str(r.get("date", "")))
    return rows

@app.get("/api/shifts/saved/export")
async def export_saved_shifts(u: str = Depends(auth), date: str = Query(None), name: str = Query(None)):
    rows = _load_saved_shifts(u)
    if date:
        rows = [r for r in rows if date in str(r.get("date", ""))]
    if name:
        nl = name.lower()
        rows = [r for r in rows if nl in str(r.get("worker_name", "")).lower()]
    rows.sort(key=lambda r: str(r.get("date", "")))

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    STATUS_COLORS = {"ok": "C6EFCE", "gap": "FFEB9C", "missing_msg": "FFCC99", "missing_excel": "FFC7CE"}
    wb = Workbook()
    ws = wb.active
    ws.title = "שעות עובדים"
    ws.sheet_view.rightToLeft = True
    headers = ["תאריך", "שם עובד", "מקום עבודה", "שעת התחלה", "שעת סיום", "שעות", "מכירות", "הערות", "סטטוס", "עובד מזוהה", "סניף מזוהה"]
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=ci, value=h)
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", start_color="D9D9D9")
    for ri, row in enumerate(rows, 2):
        vals = [row.get("date",""), row.get("worker_name",""), row.get("workplace",""),
                row.get("start_time",""), row.get("end_time",""), row.get("hours",""),
                row.get("sales",""), row.get("notes",""), row.get("status",""),
                row.get("worker_key",""), row.get("branch_key","")]
        color = STATUS_COLORS.get(row.get("status",""), "FFFFFF")
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row=ri, column=ci, value=v)
            c.fill = PatternFill("solid", start_color=color)
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        wb.save(tmp.name)
        data = Path(tmp.name).read_bytes()
    return Response(content=data,
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": "attachment; filename=worker_hours.xlsx"})

@app.post("/api/shifts/saved/upload")
async def upload_saved_shifts(u: str = Depends(auth), file: UploadFile = File(...)):
    import pandas as pd, uuid as _uuid
    contents = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(contents)
        tmp_path = Path(tmp.name)
    try:
        df = pd.read_excel(str(tmp_path), dtype=str).fillna("")
        df.columns = [str(c).strip() for c in df.columns]

        # Keyword-based column matching (order matters: more specific keywords first)
        keywords = [
            ("date",       ["תאריך", "date"]),
            ("start_time", ["התחלה", "start"]),
            ("end_time",   ["סיום", "end"]),
            ("hours",      ["שעות", "hours"]),
            ("worker_name",["שם עובד", "עובד", "worker", "name"]),
            ("workplace",  ["מקום עבודה", "מקום", "סניף", "branch", "workplace"]),
            ("sales",      ["מכירות", "sales"]),
            ("notes",      ["הערות", "notes"]),
            ("status",     ["סטטוס", "status"]),
        ]
        field_col: dict[str, str] = {}
        for field, kws in keywords:
            for col in df.columns:
                if col in field_col.values():
                    continue  # each column can only serve one field
                col_l = col.lower()
                if any(kw.lower() in col_l for kw in kws):
                    field_col[field] = col
                    break

        # Positional fallback (same order as our own export): date, worker, workplace, start, end, hours, sales, notes, status
        positional = ["date", "worker_name", "workplace", "start_time", "end_time", "hours", "sales", "notes", "status"]
        for i, field in enumerate(positional):
            if field not in field_col and i < len(df.columns) and df.columns[i] not in field_col.values():
                field_col[field] = df.columns[i]

        # Resolve worker + workplace text -> workers table / canonical branch label
        branch_map = _build_branch_map(_load_branches(u))
        resolve = _worker_resolver(u)

        def fix_date_cell(v: str) -> str:
            # Excel date cells read as str() come out "2026-07-15 00:00:00" (or with a
            # non-midnight time attached) instead of the expected "15/07/2026" — pull
            # just the date part out and reformat.
            m = re.match(r"^(\d{4})-(\d{2})-(\d{2})(?:[ T]\d{2}:\d{2}:\d{2})?$", v)
            if m:
                yr, mo, dy = m.groups()
                return f"{dy}/{mo}/{yr}"
            return v

        def fix_time_cell(v: str) -> str:
            # Excel time-only cells read as str() come out "1899-12-30 08:00:00" — pull
            # just HH:MM out.
            m = re.match(r"^(?:\d{4}-\d{2}-\d{2}[ T])?(\d{2}):(\d{2})(?::\d{2})?$", v)
            if m:
                return f"{m.group(1)}:{m.group(2)}"
            return v

        rows = []
        for _, row in df.iterrows():
            r = {"id": str(_uuid.uuid4())}
            for field in ["date", "worker_name", "workplace", "start_time", "end_time", "hours", "sales", "notes", "status"]:
                col = field_col.get(field)
                if col is not None:
                    v = str(row[col]).strip()
                    v = "" if v in ("nan", "None") else v
                    if field == "date":
                        v = fix_date_cell(v)
                    elif field in ("start_time", "end_time"):
                        v = fix_time_cell(v)
                    r[field] = v
                else:
                    r[field] = ""
            if not r.get("worker_name") and not r.get("date"):
                continue
            _enrich_shift_row(r, resolve, branch_map)
            rows.append(r)
        existing = _load_saved_shifts(u)
        existing.extend(rows)
        _save_saved_shifts(u, existing)
        return {"ok": True, "count": len(rows)}
    finally:
        tmp_path.unlink(missing_ok=True)

@app.put("/api/shifts/saved/{row_id}")
async def update_saved_shift(row_id: str, body: dict, u: str = Depends(auth)):
    rows = _load_saved_shifts(u)
    for i, r in enumerate(rows):
        if r.get("id") == row_id:
            # Merge so fields the client didn't send (e.g. worker_key/branch_key)
            # aren't wiped, then re-resolve against the current edited text
            merged = {**r, **body, "id": row_id}
            _enrich_shift_row(merged, _worker_resolver(u), _build_branch_map(_load_branches(u)), prev=r)
            rows[i] = merged
            _save_saved_shifts(u, rows)
            return {"ok": True}
    raise HTTPException(404, "שורה לא נמצאה")

@app.delete("/api/shifts/saved/{row_id}")
async def delete_saved_shift(row_id: str, u: str = Depends(auth)):
    rows = _load_saved_shifts(u)
    rows = [r for r in rows if r.get("id") != row_id]
    _save_saved_shifts(u, rows)
    return {"ok": True}


# ── Workers ───────────────────────────────────────────────────────────────────

def _workers_path(u: str) -> Path:
    return _udir(u) / "workers.json"

def _load_workers(u: str) -> list:
    return _rj(_workers_path(u), [])

def _save_workers(u: str, workers: list):
    _workers_path(u).write_text(json.dumps(workers, ensure_ascii=False, indent=2), encoding="utf-8")

@app.post("/api/workers/upload")
async def upload_workers(u: str = Depends(auth), file: UploadFile = File(...)):
    import uuid as _uuid
    import pandas as pd
    data = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        df = pd.read_excel(tmp_path, header=0)
        df.columns = [str(c).strip() for c in df.columns]

        # Map column names to internal fields — partial case-insensitive match
        keywords = {
            "full_name":      ["שם", "name"],
            "id_number":      ["ז", "id", "עובד", "מספר"],
            "nickname":       ["כינוי", "nick"],
            "manager":        ["מנהל", "manager"],
            "rank":           ["דרגה", "תפקיד", "type", "rank"],
            "sales_target":   ["יעד מכירות", "יעד", "target", "sales"],
            "arnakot_target": ["יעד ארנוק", "ארנוק"],
        }
        # Build field→column mapping
        field_col: dict[str, str] = {}
        for field, kws in keywords.items():
            for col in df.columns:
                if col in field_col.values():
                    continue  # each column serves one field
                col_l = col.lower()
                if any(kw.lower() in col_l for kw in kws) and field not in field_col:
                    field_col[field] = col
                    break

        # Positional fallback: name, id, nickname, manager, rank, sales_target
        positional = ["full_name", "id_number", "nickname", "manager", "rank", "sales_target"]
        for i, field in enumerate(positional):
            if field not in field_col and i < len(df.columns) and df.columns[i] not in field_col.values():
                field_col[field] = df.columns[i]

        def cell(row, field):
            col = field_col.get(field)
            if col is None:
                return ""
            val = row[col]
            return "" if (val is None or (isinstance(val, float) and pd.isna(val))) else str(val).strip()

        # Preserve UI-assigned salary fields across a re-upload (matched by id, then name)
        existing = _load_workers(u)
        prev_by_id   = {w.get("id_number", "").strip(): w for w in existing if w.get("id_number")}
        prev_by_name = {w.get("full_name", "").strip(): w for w in existing if w.get("full_name")}

        workers = []
        for _, row in df.iterrows():
            full_name = cell(row, "full_name")
            if not full_name:
                continue
            rank_val = cell(row, "rank").lower()
            rank = "manager" if any(x in rank_val for x in ["מנהל", "manager"]) else "worker"
            id_number = cell(row, "id_number")
            prev = prev_by_id.get(id_number) or prev_by_name.get(full_name) or {}
            workers.append({
                "id":             str(_uuid.uuid4()),
                "full_name":      full_name,
                "id_number":      id_number,
                "nickname":       cell(row, "nickname"),
                "manager":        cell(row, "manager"),
                "rank":           rank,
                "sales_target":   cell(row, "sales_target"),
                "arnakot_target": cell(row, "arnakot_target") or prev.get("arnakot_target", ""),
                "salary_model":   prev.get("salary_model", ""),
            })
        _save_workers(u, workers)
        return {"ok": True, "count": len(workers)}
    finally:
        Path(tmp_path).unlink(missing_ok=True)

@app.get("/api/workers")
async def get_workers(u: str = Depends(auth)):
    return _load_workers(u)

@app.post("/api/workers")
async def add_worker(body: dict, u: str = Depends(auth)):
    import uuid as _uuid
    workers = _load_workers(u)
    body["id"] = str(_uuid.uuid4())
    workers.append(body)
    _save_workers(u, workers)
    return {"ok": True, "worker": body}

@app.put("/api/workers/{worker_id}")
async def update_worker(worker_id: str, body: dict, u: str = Depends(auth)):
    workers = _load_workers(u)
    for i, w in enumerate(workers):
        if w.get("id") == worker_id:
            body["id"] = worker_id
            workers[i] = body
            _save_workers(u, workers)
            return {"ok": True}
    raise HTTPException(404, "עובד לא נמצא")

@app.delete("/api/workers/{worker_id}")
async def delete_worker(worker_id: str, u: str = Depends(auth)):
    workers = _load_workers(u)
    workers = [w for w in workers if w.get("id") != worker_id]
    _save_workers(u, workers)
    return {"ok": True}


# ── Branches ──────────────────────────────────────────────────────────────────

def _branches_path(u: str) -> Path:
    return _udir(u) / "branches.json"

def _load_branches(u: str) -> list:
    return _rj(_branches_path(u), [])

def _save_branches(u: str, branches: list):
    _branches_path(u).write_text(json.dumps(branches, ensure_ascii=False, indent=2), encoding="utf-8")

@app.get("/api/branches")
async def get_branches(u: str = Depends(auth)):
    # Include the canonical label so the frontend never has to re-implement
    # the normalization rules (matching.BranchMatcher owns them)
    branches = _load_branches(u)
    for b in branches:
        b["label"] = BranchMatcher.normalize_label(
            f"{(b.get('number') or '').strip()} - {(b.get('name') or '').strip()}"
        )
    return branches

@app.delete("/api/branches")
async def delete_all_branches(u: str = Depends(auth)):
    _save_branches(u, [])
    return {"ok": True}

@app.post("/api/branches/upload")
async def upload_branches(u: str = Depends(auth), file: UploadFile = File(...)):
    import uuid as _uuid
    import pandas as pd
    data = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        df = pd.read_excel(tmp_path, header=0)
        df.columns = [str(c).strip() for c in df.columns]

        keywords = {
            "number":    ["מספר", "number", "num", "מס"],
            "name":      ["שם", "name", "סניף"],
            "nicknames": ["כינוי", "nick"],
        }
        field_col: dict[str, str] = {}
        for field, kws in keywords.items():
            for col in df.columns:
                if col in field_col.values():
                    continue  # each column can only serve one field
                col_l = col.lower()
                if any(kw.lower() in col_l for kw in kws):
                    field_col[field] = col
                    break

        positional = ["number", "name", "nicknames"]
        for i, field in enumerate(positional):
            if field not in field_col and i < len(df.columns) and df.columns[i] not in field_col.values():
                field_col[field] = df.columns[i]

        def cell(row, field):
            col = field_col.get(field)
            if col is None:
                return ""
            val = row[col]
            return "" if (val is None or (isinstance(val, float) and pd.isna(val))) else str(val).strip()

        branches = _load_branches(u)
        added = 0
        for _, row in df.iterrows():
            name = cell(row, "name")
            if not name:
                continue
            nicknames_raw = cell(row, "nicknames")
            nicknames = [n.strip() for n in nicknames_raw.split(",") if n.strip()] if nicknames_raw else []
            branches.append({
                "id":        str(_uuid.uuid4()),
                "number":    cell(row, "number"),
                "name":      name,
                "nicknames": nicknames,
            })
            added += 1
        _save_branches(u, branches)
        return {"ok": True, "count": added}
    finally:
        Path(tmp_path).unlink(missing_ok=True)

@app.post("/api/branches")
async def add_branch(body: dict, u: str = Depends(auth)):
    import uuid as _uuid
    branches = _load_branches(u)
    body["id"] = str(_uuid.uuid4())
    body["nicknames"] = [n.strip() for n in (body.get("nicknames") or []) if n.strip()]
    branches.append(body)
    _save_branches(u, branches)
    return {"ok": True, "branch": body}

@app.put("/api/branches/{branch_id}")
async def update_branch(branch_id: str, body: dict, u: str = Depends(auth)):
    branches = _load_branches(u)
    for i, b in enumerate(branches):
        if b.get("id") == branch_id:
            body["id"] = branch_id
            body["nicknames"] = [n.strip() for n in (body.get("nicknames") or []) if n.strip()]
            branches[i] = body
            _save_branches(u, branches)
            return {"ok": True}
    raise HTTPException(404, "סניף לא נמצא")

@app.delete("/api/branches/{branch_id}")
async def delete_branch(branch_id: str, u: str = Depends(auth)):
    branches = _load_branches(u)
    branches = [b for b in branches if b.get("id") != branch_id]
    _save_branches(u, branches)
    return {"ok": True}

@app.get("/api/branches/discover")
async def discover_branches(u: str = Depends(auth)):
    """Distinct branch strings ('number - name') seen in sales data that aren't in the branches table yet."""
    sales = _load_sales(u)
    known = {f"{b.get('number','').strip()} - {b.get('name','').strip()}".strip(" -") for b in _load_branches(u)}
    seen = {}
    for s in sales:
        raw = (s.get("branch") or "").strip()
        if not raw or raw in known:
            continue
        seen[raw] = True
    out = []
    for raw in seen:
        if "-" in raw:
            num, _, name = raw.partition("-")
        else:
            num, name = "", raw
        out.append({"raw": raw, "number": num.strip(), "name": name.strip()})
    return out


# ── Sales ─────────────────────────────────────────────────────────────────────

def _sales_path(u: str) -> Path:
    return _udir(u) / "sales.json"

def _load_sales(u: str) -> list:
    return _rj(_sales_path(u), [])

def _save_sales(u: str, sales: list):
    _sales_path(u).write_text(json.dumps(sales, ensure_ascii=False, indent=2), encoding="utf-8")

@app.get("/api/sales")
async def get_sales(u: str = Depends(auth)):
    return _load_sales(u)

@app.post("/api/sales/upload")
async def upload_sales(u: str = Depends(auth), file: UploadFile = File(...)):
    import pandas as pd
    contents = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(contents)
        tmp_path = Path(tmp.name)
    try:
        target_sheet = "פירוט בקשות מעודכן"
        xl = pd.ExcelFile(str(tmp_path))
        sheet_name = None
        for s in xl.sheet_names:
            if s.strip() == target_sheet:
                sheet_name = s
                break
        if sheet_name is None:
            raise HTTPException(status_code=400, detail=f"גיליון '{target_sheet}' לא נמצא בקובץ")

        # Read all columns as text for easy cell access
        df = pd.read_excel(str(tmp_path), sheet_name=sheet_name, header=0, dtype=str)
        df = df.fillna("")
        # Read date column (B = index 1) separately as native datetime so Excel serial is parsed correctly
        df_dates = pd.read_excel(str(tmp_path), sheet_name=sheet_name, header=0, usecols=[1], parse_dates=[0])
        date_col = df_dates.iloc[:, 0]

        def cell(row, col_idx):
            if col_idx < len(row):
                v = str(row.iloc[col_idx]).strip()
                return "" if v in ("nan", "None") else v
            return ""

        def fmt_date(val) -> str:
            try:
                return pd.Timestamp(val).strftime("%d/%m/%Y")
            except Exception:
                return ""

        sales = []
        for i, (_, row) in enumerate(df.iterrows()):
            sale_num = cell(row, 0)   # A
            if not sale_num:
                continue
            date_val = fmt_date(date_col.iloc[i]) if i < len(date_col) else ""   # B
            branch   = cell(row, 2)   # C
            first_name = cell(row, 3) # D
            last_name  = cell(row, 4) # E
            standing_order_raw = cell(row, 10)  # K
            standing_order = standing_order_raw == "מולא"
            revolving_l    = cell(row, 11) == "1"  # L
            revolving_m    = cell(row, 12) == "1"  # M
            revolving_h    = cell(row, 13) == "1"  # N
            revolving_xl   = cell(row, 14) == "1"  # O
            status_raw = cell(row, 15)              # P
            approved = status_raw != "תעודה מזהה לא בתוקף" and status_raw != ""

            sales.append({
                "sale_number":    sale_num,
                "date":           date_val,
                "branch":         branch,
                "first_name":     first_name,
                "last_name":      last_name,
                "standing_order": standing_order,
                "revolving_1500": revolving_l,
                "revolving_2500": revolving_m,
                "revolving_4000": revolving_h,
                "revolving_4001": revolving_xl,
                "approved":       approved,
                "status_raw":     status_raw,
            })
        _save_sales(u, sales)
        return {"ok": True, "count": len(sales)}
    finally:
        tmp_path.unlink(missing_ok=True)


# ── Report colors ─────────────────────────────────────────────────────────────

def _report_colors_path(u: str) -> Path:
    return _udir(u) / "report_colors.json"

def _load_report_colors(u: str) -> dict:
    return _rj(_report_colors_path(u), {"managers": {}, "sum_color": "#bdd7ee", "grand_color": "#ffd966"})

@app.get("/api/report/colors")
async def get_report_colors(u: str = Depends(auth)):
    return _load_report_colors(u)

@app.put("/api/report/colors")
async def put_report_colors(body: dict, u: str = Depends(auth)):
    colors = {
        "managers":    body.get("managers", {}) or {},
        "sum_color":   body.get("sum_color", "#bdd7ee"),
        "grand_color": body.get("grand_color", "#ffd966"),
    }
    _report_colors_path(u).write_text(json.dumps(colors, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True}


# ── Arnakot (ארנוקים) ─────────────────────────────────────────────────────────

def _arnakot_path(u: str) -> Path:
    return _udir(u) / "arnakot.json"

def _load_arnakot(u: str) -> list:
    return _rj(_arnakot_path(u), [])

def _save_arnakot(u: str, data: list):
    _arnakot_path(u).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

@app.get("/api/arnakot")
async def get_arnakot(u: str = Depends(auth)):
    return _load_arnakot(u)

@app.post("/api/arnakot/upload")
async def upload_arnakot(u: str = Depends(auth), file: UploadFile = File(...)):
    import pandas as pd
    contents = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(contents)
        tmp_path = Path(tmp.name)
    try:
        target_sheet = "מפורט"
        xl = pd.ExcelFile(str(tmp_path))
        sheet_name = None
        for s in xl.sheet_names:
            if s.strip() == target_sheet:
                sheet_name = s
                break
        if sheet_name is None:
            raise HTTPException(status_code=400, detail=f"גיליון '{target_sheet}' לא נמצא בקובץ")

        df = pd.read_excel(str(tmp_path), sheet_name=sheet_name, header=0, dtype=str)
        df = df.fillna("")
        # Read date column G (index 6) as native datetime to avoid format ambiguity
        df_dates = pd.read_excel(str(tmp_path), sheet_name=sheet_name, header=0, usecols=[6], parse_dates=[0])
        date_col = df_dates.iloc[:, 0]

        def fmt_date(val) -> str:
            try:
                return pd.Timestamp(val).strftime("%d/%m/%Y")
            except Exception:
                return ""

        workers = _load_workers(u)
        known_names = {w.get("full_name", "").strip() for w in workers if w.get("full_name")}

        records = []
        unmatched = set()
        for i, (_, row) in enumerate(df.iterrows()):
            if i >= len(date_col):
                continue
            name = str(row.iloc[4]).strip() if len(row) > 4 else ""  # column E
            if not name or name in ("nan", "None", ""):
                continue
            date_val = fmt_date(date_col.iloc[i])
            if not date_val:
                continue
            if name not in known_names:
                unmatched.add(name)
            records.append({"name": name, "date": date_val})

        _save_arnakot(u, records)
        return {"ok": True, "count": len(records), "unmatched": sorted(unmatched)}
    finally:
        tmp_path.unlink(missing_ok=True)


def _build_report_tables(u: str, month: str = None):
    """Builds monthly + daily report data as neutral structures.
    Returns (report_month, tables) where each table is
    {"title": str|None, "headers": [...], "rows": [{"cells": [...], "fill": hex|None, "bold": bool}]}"""
    from datetime import date as _date
    import calendar as _calendar

    workers     = _load_workers(u)
    shifts_all  = _load_saved_shifts(u)
    sales_all   = _load_sales(u)
    arnakot_all = _load_arnakot(u)

    # Arnakot stats (filtered to month)
    arnakot_stats: dict = {}

    # Determine month to report on
    all_dates = (
        [str(s.get("date","")) for s in shifts_all] +
        [str(s.get("date","")) for s in sales_all]
    )
    all_yms = sorted({_date_to_ym(d) for d in all_dates if _date_to_ym(d)}, reverse=True)
    report_month = month or (all_yms[0] if all_yms else "")

    def in_month(d: str) -> bool:
        return bool(report_month) and _date_to_ym(str(d)) == report_month

    # Work day calculation
    proj_total = proj_done = 0.0
    until_iso = ""
    if report_month:
        try:
            yr, mo = int(report_month[:4]), int(report_month[5:7])
            holidays_list = _load_holidays_cached(yr) or _fetch_and_build_holidays(yr)
            holiday_map = {h["date"]: h["type"] for h in holidays_list}

            # Latest date in data for this month
            month_dates = [_date_to_ym(str(d)) == report_month and d for d in all_dates]
            def d_to_iso(d: str) -> str:
                d = d.strip()
                if len(d) >= 10 and d[2] == "/" and d[5] == "/":
                    return f"{d[6:10]}-{d[3:5]}-{d[0:2]}"
                return d
            iso_dates = [d_to_iso(str(s.get("date",""))) for s in shifts_all + sales_all if in_month(str(s.get("date","")))]
            until_iso = max(iso_dates) if iso_dates else ""

            days_in = _calendar.monthrange(yr, mo)[1]
            for day in range(1, days_in + 1):
                d = _date(yr, mo, day)
                dow = d.weekday()  # 0=Mon … 5=Sat … 6=Sun
                if dow == 5:  # Saturday
                    continue
                d_iso = d.isoformat()
                htype = holiday_map.get(d_iso)
                if htype == "off":
                    val = 0.0
                elif htype == "half" or dow == 4:  # Friday = 4 in Python (Mon=0)
                    val = 0.5
                else:
                    val = 1.0
                proj_total += val
                if until_iso and d_iso <= until_iso:
                    proj_done += val
        except Exception:
            pass

    # Shift stats (filtered to month) — keyed by the resolved worker when available
    shift_stats: dict = {}
    for s in shifts_all:
        if not in_month(str(s.get("date",""))):
            continue
        name = (s.get("worker_key") or s.get("worker_name") or "").strip()
        if not name:
            continue
        if name not in shift_stats:
            shift_stats[name] = {"dates": set(), "hours": 0.0}
        if s.get("date"):
            shift_stats[name]["dates"].add(s["date"])
        shift_stats[name]["hours"] += float(s.get("hours") or 0)

    # Sales stats (filtered to month)
    def init_sale():
        return {"total": 0, "rev1500": 0, "rev2500": 0, "rev4000": 0, "so": 0, "issued": 0}
    sale_stats: dict = {}
    for s in sales_all:
        if not in_month(str(s.get("date",""))):
            continue
        name = ((s.get("first_name") or "") + " " + (s.get("last_name") or "")).strip()
        if not name:
            continue
        if name not in sale_stats:
            sale_stats[name] = init_sale()
        st = sale_stats[name]
        if s.get("approved"):       st["total"]  += 1
        if s.get("revolving_1500"): st["rev1500"] += 1
        if s.get("revolving_2500"): st["rev2500"] += 1
        if s.get("revolving_4000"): st["rev4000"] += 1
        if s.get("standing_order"): st["so"]      += 1
        if (s.get("status_raw") or "").strip() == "הונפק": st["issued"] += 1

    # Arnakot filtered to month
    for a in arnakot_all:
        if not in_month(str(a.get("date",""))):
            continue
        name = (a.get("name") or "").strip()
        if not name:
            continue
        arnakot_stats[name] = arnakot_stats.get(name, 0) + 1

    # Colors config
    colors_cfg = _load_report_colors(u)

    def hex6(h, default=None):
        h = (h or "").lstrip("#").strip()
        return h.upper() if len(h) == 6 else default

    sum_hex   = hex6(colors_cfg.get("sum_color"), "BDD7EE")
    grand_hex = hex6(colors_cfg.get("grand_color"), "FFD966")
    mgr_hexes = {name: hex6(c) for name, c in (colors_cfg.get("managers") or {}).items()}

    def sort_key(w):
        try:
            return (0, float(str(w.get("sort_order", "")).strip()))
        except (ValueError, TypeError):
            return (1, 0.0)

    # Group by manager
    manager_map: dict = {}
    no_manager: list = []
    for w in workers:
        if w.get("rank") == "manager":
            nm = w.get("full_name", "")
            if nm not in manager_map:
                manager_map[nm] = {"mgr": w, "workers": []}
            else:
                manager_map[nm]["mgr"] = w
    for w in workers:
        if w.get("rank") == "manager":
            continue
        mgr = (w.get("manager") or "").strip()
        if mgr and mgr in manager_map:
            manager_map[mgr]["workers"].append(w)
        elif mgr:
            if mgr not in manager_map:
                manager_map[mgr] = {"mgr": None, "workers": []}
            manager_map[mgr]["workers"].append(w)
        else:
            no_manager.append(w)

    # Sort workers inside each group by מסד; order groups by the manager's own מסד
    for group in manager_map.values():
        group["workers"].sort(key=sort_key)
    ordered_groups = sorted(
        manager_map.items(),
        key=lambda kv: sort_key(kv[1]["mgr"]) if kv[1]["mgr"] else (1, 0.0),
    )
    no_manager.sort(key=sort_key)

    HEADERS = [
        "ת.ז", "שם עובד", "מנהל", "משמרות", "שעות",
        "סה\"כ מכירות", "אשראי מתגלגל עד 1500", "אשראי מתגלגל 1501-2500", "אשראי מתגלגל 2501-4000",
        "סה\"כ הו\"ק", "סה\"כ ארנוקים", "ממוצע מכירות/שעה",
        "צפי שעות", "יעד", "צפי מכירות", "% הגעה ליעד", "הנפקות",
    ]

    monthly_rows: list = []
    current_rows = monthly_rows

    def fmt(n):
        if n is None or n == "": return ""
        try:
            f = float(n)
            return round(f, 2) if f % 1 else int(f)
        except Exception:
            return n

    def proj_val(current, done, total):
        if done and total:
            return fmt(current / done * total)
        return ""

    def z(v):
        # Replace blank/None numeric cells with 0
        return 0 if v in ("", None) else v

    def worker_data(w, is_mgr):
        name     = w.get("full_name", "")
        ss       = shift_stats.get(name, {"dates": set(), "hours": 0.0})
        sa       = sale_stats.get(name, init_sale())
        shifts_n = len(ss["dates"])
        hours_n  = ss["hours"]
        avg      = fmt(sa["total"] / hours_n) if hours_n else 0
        p_hours  = z(proj_val(hours_n, proj_done, proj_total))
        p_sales  = z(proj_val(sa["total"], proj_done, proj_total))
        target   = w.get("sales_target", "")
        try:
            pct = fmt(float(p_sales) / float(target) * 100) if float(target or 0) else 0
        except (ValueError, TypeError):
            pct = 0
        return [
            w.get("id_number", ""), name,
            "" if is_mgr else (w.get("manager") or ""),
            shifts_n, z(fmt(hours_n)),
            sa["total"], sa["rev1500"], sa["rev2500"], sa["rev4000"],
            sa["so"], arnakot_stats.get(name, 0), avg,
            p_hours, z(target), p_sales, pct, sa["issued"],
        ]

    def sum_data(label, name_list, worker_only_names):
        shifts_n = hours_n = total = rev1500 = rev2500 = rev4000 = so = issued = arnakot_sum = 0
        for name in name_list:
            ss = shift_stats.get(name, {"dates": set(), "hours": 0.0})
            sa = sale_stats.get(name, init_sale())
            shifts_n    += len(ss["dates"])
            hours_n     += ss["hours"]
            total       += sa["total"]
            rev1500     += sa["rev1500"]
            rev2500     += sa["rev2500"]
            rev4000     += sa["rev4000"]
            so          += sa["so"]
            issued      += sa["issued"]
            arnakot_sum += arnakot_stats.get(name, 0)
        # יעד and % הגעה: workers only (exclude the manager's own data)
        target_sum = 0.0
        wtotal = 0
        for name in worker_only_names:
            w = next((x for x in workers if x.get("full_name") == name), None)
            try:
                target_sum += float(w.get("sales_target") or 0) if w else 0
            except (ValueError, TypeError):
                pass
            wtotal += sale_stats.get(name, init_sale())["total"]
        p_sales_w = proj_val(wtotal, proj_done, proj_total)
        try:
            pct = fmt(float(p_sales_w or 0) / target_sum * 100) if target_sum else 0
        except (ValueError, TypeError):
            pct = 0
        avg     = fmt(total / hours_n) if hours_n else 0
        p_hours = z(proj_val(hours_n, proj_done, proj_total))
        p_sales = z(proj_val(total, proj_done, proj_total))
        return [
            label, label, "",
            shifts_n, z(fmt(hours_n)),
            total, rev1500, rev2500, rev4000,
            so, arnakot_sum, avg,
            p_hours, fmt(target_sum), p_sales, pct, issued,
        ]

    def append_row(row_data, fill_hex=None, bold=False):
        current_rows.append({"cells": row_data, "fill": fill_hex, "bold": bold})

    grand_names = []          # everyone, for operational sums
    grand_worker_names = []   # workers only, for target/% columns

    for mgr_name, group in ordered_groups:
        m_hex = mgr_hexes.get(mgr_name)
        for w in group["workers"]:
            append_row(worker_data(w, False), fill_hex=m_hex)
        if group["mgr"]:
            append_row(worker_data(group["mgr"], True), fill_hex=m_hex, bold=True)
        worker_names = [w.get("full_name", "") for w in group["workers"]]
        all_names = list(worker_names)
        if group["mgr"]:
            all_names.append(group["mgr"].get("full_name", ""))
        append_row(sum_data(f"סה\"כ {mgr_name}", all_names, worker_names), fill_hex=sum_hex, bold=True)
        grand_names.extend(all_names)
        grand_worker_names.extend(worker_names)

    for w in no_manager:
        append_row(worker_data(w, False))
        grand_names.append(w.get("full_name", ""))
        grand_worker_names.append(w.get("full_name", ""))

    if grand_names:
        append_row(sum_data("סה\"כ כללי", grand_names, grand_worker_names), fill_hex=grand_hex, bold=True)

    # ── Daily report section (latest data date) ──
    daily_table = None
    if until_iso:
        daily_date = f"{until_iso[8:10]}/{until_iso[5:7]}/{until_iso[0:4]}"

        d_shift: dict = {}
        for s in shifts_all:
            if str(s.get("date", "")) != daily_date: continue
            nm = (s.get("worker_key") or s.get("worker_name") or "").strip()
            if not nm: continue
            d_shift[nm] = d_shift.get(nm, 0.0) + float(s.get("hours") or 0)

        d_sale: dict = {}
        for s in sales_all:
            if str(s.get("date", "")) != daily_date: continue
            nm = ((s.get("first_name") or "") + " " + (s.get("last_name") or "")).strip()
            if not nm: continue
            st = d_sale.setdefault(nm, {"total": 0, "rev": 0, "so": 0, "issued": 0})
            if s.get("approved"): st["total"] += 1
            if s.get("revolving_1500") or s.get("revolving_2500") or s.get("revolving_4000") or s.get("revolving_4001"):
                st["rev"] += 1
            if s.get("standing_order"): st["so"] += 1
            if (s.get("status_raw") or "").strip() == "הונפק": st["issued"] += 1

        d_arnak: dict = {}
        for a in arnakot_all:
            if str(a.get("date", "")) != daily_date: continue
            nm = (a.get("name") or "").strip()
            if not nm: continue
            d_arnak[nm] = d_arnak.get(nm, 0) + 1
        arnak_updated = bool(d_arnak)

        DAILY_HEADERS = [
            "ת.ז", "שם עובד", "מנהל", "שעות", "סה\"כ מכירות", "אשראי מתגלגל",
            "הו\"ק", "ארנוקים עודכן" if arnak_updated else "ארנוקים לא מעודכן",
            "ממוצע מכירות/שעה", "הנפקות",
        ]

        daily_rows: list = []
        current_rows = daily_rows

        def daily_worker_row(w):
            nm = w.get("full_name", "")
            hr = d_shift.get(nm, 0.0)
            sa = d_sale.get(nm, {"total": 0, "rev": 0, "so": 0, "issued": 0})
            avg = fmt(sa["total"] / hr) if hr else 0
            return [
                w.get("id_number", ""), nm, w.get("manager", "") or "",
                z(fmt(hr)), sa["total"], sa["rev"], sa["so"],
                d_arnak.get(nm, 0), avg, sa["issued"],
            ]

        def daily_sum_row(label, name_list):
            hr = total = rev = so = arnak = issued = 0
            for nm in name_list:
                hr    += d_shift.get(nm, 0.0)
                sa     = d_sale.get(nm, {"total": 0, "rev": 0, "so": 0, "issued": 0})
                total += sa["total"]; rev += sa["rev"]; so += sa["so"]; issued += sa["issued"]
                arnak += d_arnak.get(nm, 0)
            avg = fmt(total / hr) if hr else 0
            return [label, label, "", z(fmt(hr)), total, rev, so, arnak, avg, issued]

        def daily_active(nm):
            return nm in d_shift or nm in d_sale or nm in d_arnak

        for mgr_name, group in ordered_groups:
            m_hex = mgr_hexes.get(mgr_name)
            block_names = []
            for w in group["workers"]:
                if daily_active(w.get("full_name", "")):
                    append_row(daily_worker_row(w), fill_hex=m_hex)
                block_names.append(w.get("full_name", ""))
            if group["mgr"]:
                if daily_active(group["mgr"].get("full_name", "")):
                    append_row(daily_worker_row(group["mgr"]), fill_hex=m_hex, bold=True)
                block_names.append(group["mgr"].get("full_name", ""))
            append_row(daily_sum_row(f"סה\"כ {mgr_name}", block_names), fill_hex=sum_hex, bold=True)

        d_grand = []
        for _, group in ordered_groups:
            d_grand.extend(w.get("full_name", "") for w in group["workers"])
            if group["mgr"]: d_grand.append(group["mgr"].get("full_name", ""))
        for w in no_manager:
            if daily_active(w.get("full_name", "")):
                append_row(daily_worker_row(w))
            d_grand.append(w.get("full_name", ""))
        if d_grand:
            append_row(daily_sum_row("סה\"כ כללי", d_grand), fill_hex=grand_hex, bold=True)

        daily_table = {"title": f"דוח יומי — {daily_date}", "headers": DAILY_HEADERS, "rows": daily_rows}

    tables = [{"title": None, "headers": HEADERS, "rows": monthly_rows}]
    if daily_table:
        tables.append(daily_table)
    return report_month, tables


def _build_branch_map(branches: list) -> dict:
    """nickname (normalized) -> canonical branch label. Thin wrapper over the
    shared BranchMatcher so existing callers keep working."""
    return BranchMatcher(branches).as_dict()


def _normalize_branch_label(text: str) -> str:
    return BranchMatcher.normalize_label(text)


def _d_to_iso(d: str) -> str:
    d = (d or "").strip()
    if len(d) >= 10 and d[2] == "/" and d[5] == "/":
        return f"{d[6:10]}-{d[3:5]}-{d[0:2]}"
    return d

def _branch_report_for_date(u: str, branches: list, sales_all: list, shifts_all: list,
                            resolve, report_date: str):
    """Branch rows + sales-without-hours mismatches for one date."""
    def in_day(d: str) -> bool:
        return bool(report_date) and str(d).strip() == report_date

    in_month = in_day  # kept as a local alias so the block below reads unchanged

    branch_label = {}
    for b in branches:
        label = _normalize_branch_label(f"{b.get('number','').strip()} - {b.get('name','').strip()}")
        branch_label[label] = b

    # Sales per branch (every row counts, no approval filter)
    sales_count: dict = {}
    sales_by_worker_branch: dict = {}  # branch_label -> {worker name: [sale dicts]}
    for s in sales_all:
        if not in_month(str(s.get("date",""))):
            continue
        raw = _normalize_branch_label((s.get("branch") or "").strip())
        if not raw:
            continue
        sales_count[raw] = sales_count.get(raw, 0) + 1
        name = ((s.get("first_name") or "") + " " + (s.get("last_name") or "")).strip()
        if name:
            sales_by_worker_branch.setdefault(raw, {}).setdefault(name, []).append({
                "sale_number": s.get("sale_number", ""),
                "status_raw":  s.get("status_raw", ""),
            })

    # Hours per resolved branch; lines with no branch link go to "לא משויך"
    hours_by_branch: dict = {}
    unassigned_hours = 0.0
    for s in shifts_all:
        if not in_month(str(s.get("date",""))):
            continue
        bkey = _normalize_branch_label((s.get("branch_key") or "").strip())
        if not bkey:
            unassigned_hours += float(s.get("hours") or 0)
            continue
        hours_by_branch[bkey] = hours_by_branch.get(bkey, 0.0) + float(s.get("hours") or 0)

    all_branch_keys = sorted(set(sales_count) | set(hours_by_branch))
    rows = []
    for key in all_branch_keys:
        b = branch_label.get(key, {})
        rows.append({
            "date":   report_date,
            "number": b.get("number", key.split(" - ")[0] if " - " in key else ""),
            "name":   b.get("name", key.split(" - ", 1)[1] if " - " in key else key),
            "label":  key,
            "hours":  round(hours_by_branch.get(key, 0.0), 2),
            "sales":  sales_count.get(key, 0),
        })
    if unassigned_hours:
        rows.append({
            "date": report_date,
            "number": "", "name": "⚠️ לא משויך לסניף", "label": "",
            "hours": round(unassigned_hours, 2), "sales": 0,
        })

    # Mismatch: branches with sales but zero hours logged this day
    mismatches = []
    for key in all_branch_keys:
        if sales_count.get(key, 0) > 0 and hours_by_branch.get(key, 0.0) == 0:
            worker_details = []
            for name in sorted(sales_by_worker_branch.get(key, {})):
                full_name = resolve(name)
                shift_texts = []
                for s in shifts_all:
                    if not in_month(str(s.get("date",""))):
                        continue
                    s_names = {(s.get("worker_name") or "").strip(), (s.get("worker_key") or "").strip()} - {""}
                    candidates = {name, full_name} - {""}
                    if not (candidates & s_names):
                        continue
                    shift_texts.append({
                        "date": s.get("date", ""),
                        "workplace": s.get("workplace", ""),
                        "branch_key": s.get("branch_key", ""),
                    })
                worker_details.append({
                    "worker_name": name,
                    "full_name":   full_name or name,
                    "sales":       sales_by_worker_branch.get(key, {}).get(name, []),
                    "shifts":      shift_texts,
                })
            mismatches.append({"branch_label": key, "workers": worker_details})

    return rows, mismatches


def _branch_report_all_dates(sales_all: list, shifts_all: list) -> list:
    """Every distinct date present in sales/shifts, oldest first."""
    seen = {}
    for s in (sales_all + shifts_all):
        d = str(s.get("date", "")).strip()
        if d:
            seen[d] = _d_to_iso(d)
    return [d for d, _iso in sorted(seen.items(), key=lambda kv: kv[1])]


def _branch_report_payload(u: str, date: str = None, all_days: bool = False):
    branches   = _load_branches(u)
    sales_all  = _load_sales(u)
    shifts_all = _load_saved_shifts(u)
    resolve    = _worker_resolver(u)
    dates      = _branch_report_all_dates(sales_all, shifts_all)

    if all_days:
        days = []
        for d in dates:
            rows, mism = _branch_report_for_date(u, branches, sales_all, shifts_all, resolve, d)
            if rows:
                days.append({"date": d, "rows": rows, "mismatches": mism})
        return {"all_days": True, "days": days}

    report_date = date or (dates[-1] if dates else "")
    rows, mism = _branch_report_for_date(u, branches, sales_all, shifts_all, resolve, report_date)
    return {"all_days": False, "date": report_date, "rows": rows, "mismatches": mism}


@app.get("/api/report/branches")
async def report_branches(u: str = Depends(auth), date: str = Query(None),
                          all_days: bool = Query(False)):
    return _branch_report_payload(u, date, all_days)


@app.get("/api/report/branches/export")
async def export_branch_report(u: str = Depends(auth), date: str = Query(None),
                               all_days: bool = Query(False)):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    data = _branch_report_payload(u, date, all_days)
    wb = Workbook()
    ws = wb.active
    ws.title = "דוח סניפים"
    ws.sheet_view.rightToLeft = True

    hdr_fill = PatternFill("solid", fgColor="2F5496")
    hdr_font = Font(bold=True, color="FFFFFF", size=10)
    center   = Alignment(horizontal="center", vertical="center")
    thin     = Side(style="thin", color="B0B0B0")
    border   = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.append(["תאריך", "מספר", "שם סניף", "שעות", "מכירות"])
    for c in ws[1]:
        c.font = hdr_font; c.fill = hdr_fill; c.alignment = center; c.border = border

    groups = data["days"] if data.get("all_days") else [{"date": data.get("date", ""), "rows": data.get("rows", [])}]
    for g in groups:
        for r in g["rows"]:
            ws.append([r.get("date", g["date"]), r.get("number", ""), r.get("name", ""),
                       r.get("hours", 0), r.get("sales", 0)])
            for c in ws[ws.max_row]:
                c.alignment = center; c.border = border

    for col in ws.columns:
        w = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(w + 4, 26)

    out = Path(tempfile.mktemp(suffix=".xlsx"))
    wb.save(str(out))
    payload = out.read_bytes()
    out.unlink(missing_ok=True)
    suffix = "all_days" if data.get("all_days") else (data.get("date", "") or "report").replace("/", "-")
    return Response(
        content=payload,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="branch_report_{suffix}.xlsx"'},
    )


def _build_report_xlsx(u: str, month: str = None):
    """Returns (report_month, xlsx_bytes) for the full sales report."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    report_month, tables = _build_report_tables(u, month)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "דוח מלא"
    ws.sheet_view.rightToLeft = True

    hdr_fill   = PatternFill("solid", fgColor="2F5496")
    hdr_font   = Font(bold=True, color="FFFFFF", size=10)
    bold_font  = Font(bold=True, size=10)
    norm_font  = Font(size=10)
    center     = Alignment(horizontal="center", vertical="center")
    thin       = Side(style="thin", color="808080")
    border     = Border(left=thin, right=thin, top=thin, bottom=thin)

    for t_idx, table in enumerate(tables):
        if t_idx > 0:
            ws.append([])
        if table["title"]:
            ws.append([table["title"]])
            ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=12)
        ws.append(table["headers"])
        for cell in ws[ws.max_row]:
            cell.font = hdr_font; cell.fill = hdr_fill
            cell.alignment = center; cell.border = border
        for r in table["rows"]:
            ws.append(r["cells"])
            fill = PatternFill("solid", fgColor=r["fill"]) if r["fill"] else None
            for cell in ws[ws.max_row]:
                cell.font = bold_font if r["bold"] else norm_font
                if fill: cell.fill = fill
                cell.alignment = center
                cell.border = border

    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 28)

    out = Path(tempfile.mktemp(suffix=".xlsx"))
    wb.save(str(out))
    data = out.read_bytes()
    out.unlink(missing_ok=True)
    return report_month, data


@app.get("/api/report/export")
async def export_report(u: str = Depends(auth), month: str = Query(None)):
    report_month, data = _build_report_xlsx(u, month)
    fname = f"worker_report_{report_month or 'all'}.xlsx"
    from fastapi.responses import Response
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


def _build_report_pdf(u: str, month: str = None):
    """Returns (report_month, pdf_bytes) for the full sales report."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from bidi.algorithm import get_display
    import io

    report_month, tables = _build_report_tables(u, month)

    font_path = BASE_DIR / "fonts" / "DejaVuSans.ttf"
    font_bold_path = BASE_DIR / "fonts" / "DejaVuSans-Bold.ttf"
    if "DejaVu" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVu", str(font_path)))
        pdfmetrics.registerFont(TTFont("DejaVu-Bold", str(font_bold_path)))

    def heb(text):
        # bidi-reorder Hebrew for correct PDF display
        return get_display(str(text))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(A4),
        rightMargin=8*mm, leftMargin=8*mm, topMargin=10*mm, bottomMargin=10*mm,
    )

    title_style = ParagraphStyle("title", fontName="DejaVu-Bold", fontSize=13, alignment=2)
    elements = []

    month_label = report_month.split("-")[::-1] if report_month else []
    month_title = f"דוח מלא — {'/'.join(month_label)}" if month_label else "דוח מלא"
    elements.append(Paragraph(heb(month_title), title_style))
    elements.append(Spacer(1, 4*mm))

    def hex_to_color(h):
        try:
            return rl_colors.HexColor(f"#{h}")
        except Exception:
            return None

    for table in tables:
        if table["title"]:
            elements.append(Spacer(1, 6*mm))
            elements.append(Paragraph(heb(table["title"]), title_style))
            elements.append(Spacer(1, 3*mm))

        # RTL: reverse column order so the first logical column is rightmost
        header = [heb(h) for h in reversed(table["headers"])]
        data = [header]
        style_cmds = [
            ("FONTNAME",  (0, 0), (-1, 0), "DejaVu-Bold"),
            ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#2F5496")),
            ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
            ("FONTNAME",  (0, 1), (-1, -1), "DejaVu"),
            ("FONTSIZE",  (0, 0), (-1, -1), 6.5),
            ("ALIGN",     (0, 0), (-1, -1), "CENTER"),
            ("VALIGN",    (0, 0), (-1, -1), "MIDDLE"),
            ("GRID",      (0, 0), (-1, -1), 0.4, rl_colors.grey),
            ("TOPPADDING",    (0, 0), (-1, -1), 1.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5),
            ("LEFTPADDING",   (0, 0), (-1, -1), 2),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 2),
        ]
        for i, r in enumerate(table["rows"], start=1):
            data.append([heb(c) for c in reversed(r["cells"])])
            if r["fill"]:
                c = hex_to_color(r["fill"])
                if c:
                    style_cmds.append(("BACKGROUND", (0, i), (-1, i), c))
            if r["bold"]:
                style_cmds.append(("FONTNAME", (0, i), (-1, i), "DejaVu-Bold"))

        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle(style_cmds))
        elements.append(t)

    doc.build(elements)
    return report_month, buf.getvalue()


@app.get("/api/report/export/pdf")
async def export_report_pdf(u: str = Depends(auth), month: str = Query(None)):
    report_month, pdf_data = _build_report_pdf(u, month)
    fname = f"worker_report_{report_month or 'all'}.pdf"
    from fastapi.responses import Response
    return Response(
        content=pdf_data,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ── Report email (sales report as PDF + Excel via the Gmail HTTPS API) ─────────
# Railway blocks outbound SMTP, so sending goes through the Gmail API from one
# fixed sender account connected once via OAuth (see /auth/report-gmail/connect).

def _report_email_path(u: str) -> Path:
    return _udir(u) / "report_email.json"

# {תאריך} in the subject/body is replaced at send time with the last date in the report
DATE_TOKEN = "{תאריך}"
_DEFAULT_REPORT_SUBJECT = "דוח מכירות — לייף סטייל " + DATE_TOKEN
_DEFAULT_REPORT_BODY = (
    "שלום,\n"
    "מצורף דוח המכירות המלא (PDF ו-Excel) נכון לתאריך " + DATE_TOKEN + ".\n\n"
    "Electra Target Tools"
)

def _load_report_email(u: str) -> dict:
    d = _rj(_report_email_path(u), {})
    # Migrate the old single-recipient shape
    recipients = d.get("recipients")
    if recipients is None:
        one = (d.get("recipient") or "").strip()
        recipients = [one] if one else []
    return {
        "recipients": [r for r in recipients if r],
        "subject": d.get("subject") if d.get("subject") is not None else _DEFAULT_REPORT_SUBJECT,
        "body":    d.get("body")    if d.get("body")    is not None else _DEFAULT_REPORT_BODY,
    }

def _save_report_email(u: str, cfg: dict):
    _wj(_report_email_path(u), cfg)

def _report_last_date(u: str, month: str = None) -> str:
    """Latest data date (DD/MM/YYYY) in the report — the month's last date, or
    the newest date overall when no month is given."""
    all_dates = [str(s.get("date", "")) for s in (_load_saved_shifts(u) + _load_sales(u)) if s.get("date")]
    all_dates = [d for d in all_dates if _d_to_iso(d)]
    if not all_dates:
        return ""
    if not month:
        month = max(_date_to_ym(d) for d in all_dates if _date_to_ym(d))
    dates = [d for d in all_dates if _date_to_ym(d) == month]
    return max(dates, key=_d_to_iso) if dates else ""

@app.get("/api/report/email/settings")
async def report_email_settings(u: str = Depends(auth), month: str = Query(None)):
    cfg = _load_report_email(u)
    tok = _get_valid_gmail_tokens(REPORT_SENDER_KEY)
    return {
        **cfg,
        "configured": bool(tok),
        "sender": _report_sender_email(),
        "date_token": DATE_TOKEN,
        "last_date": _report_last_date(u, month),
    }

@app.put("/api/report/email/settings")
async def save_report_email_settings(body: dict, u: str = Depends(auth)):
    recips = body.get("recipients")
    if recips is None and body.get("recipient") is not None:  # tolerate single
        recips = [body.get("recipient")]
    cfg = {
        "recipients": [r.strip() for r in (recips or []) if r and r.strip()],
        "subject": (body.get("subject") if body.get("subject") is not None else _DEFAULT_REPORT_SUBJECT),
        "body":    (body.get("body")    if body.get("body")    is not None else _DEFAULT_REPORT_BODY),
    }
    _save_report_email(u, cfg)
    return {"ok": True}

@app.post("/api/report/email/test")
async def test_report_email(u: str = Depends(auth)):
    """Confirm the connected sender account still works, without sending."""
    tok = _get_valid_gmail_tokens(REPORT_SENDER_KEY)
    if not tok:
        raise HTTPException(400, "חשבון Gmail לשליחה אינו מחובר")
    import urllib.request, json as _json
    try:
        req = urllib.request.Request(
            "https://gmail.googleapis.com/gmail/v1/users/me/profile",
            headers={"Authorization": f"Bearer {tok['access_token']}"})
        with urllib.request.urlopen(req, timeout=15) as r:
            email = _json.loads(r.read()).get("emailAddress", "")
    except Exception as e:
        raise HTTPException(502, f"בדיקת החיבור נכשלה: {e}")
    return {"ok": True, "sender": email or _report_sender_email()}

@app.post("/api/report/email/send")
async def send_report_email(body: dict, u: str = Depends(auth)):
    import base64
    from emailer import build_email_message

    tok = _get_valid_gmail_tokens(REPORT_SENDER_KEY)
    if not tok:
        raise HTTPException(400, "חשבון Gmail לשליחה אינו מחובר")
    sender = _report_sender_email() or "me"

    cfg = _load_report_email(u)
    # Body may pass overrides (recipients/subject/body); fall back to saved config
    recipients = body.get("recipients")
    if recipients is None:
        recipients = cfg["recipients"]
    recipients = [r.strip() for r in recipients if r and r.strip()]
    if not recipients:
        raise HTTPException(400, "לא הוזנו נמענים")
    bad = [r for r in recipients if "@" not in r]
    if bad:
        raise HTTPException(400, f"כתובת מייל לא תקינה: {bad[0]}")

    month = body.get("month")
    report_month, xlsx = _build_report_xlsx(u, month)
    _rm, pdf = _build_report_pdf(u, month)

    last_date = _report_last_date(u, month)
    subj_tmpl = body.get("subject") if body.get("subject") is not None else cfg["subject"]
    body_tmpl = body.get("body")    if body.get("body")    is not None else cfg["body"]
    subject   = (subj_tmpl or "").replace(DATE_TOKEN, last_date)
    body_text = (body_tmpl or "").replace(DATE_TOKEN, last_date)

    import html as _html
    body_html = _html.escape(body_text).replace("\n", "<br>")
    html = f"""<html><body style="background:#f5f5f5;padding:20px">
      <div style="font-family:Arial,sans-serif;direction:rtl;text-align:right;color:#1a1a2e;font-size:14px;line-height:1.6">
        {body_html}
      </div>
    </body></html>"""

    fbase = f"sales_report_{report_month or 'all'}"
    msg = build_email_message(sender, recipients, subject, html, [
        (f"{fbase}.pdf",  pdf,  "pdf"),
        (f"{fbase}.xlsx", xlsx, "vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    ])
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    try:
        _gmail_api_send(tok["access_token"], raw)
    except Exception as e:
        # An expired access token surfaces as 401 — refresh once and retry
        import urllib.error
        if isinstance(e, urllib.error.HTTPError) and e.code == 401:
            refreshed = _refresh_gmail_tokens(REPORT_SENDER_KEY, tok)
            if not refreshed:
                raise HTTPException(400, "חיבור ה-Gmail פג, יש להתחבר מחדש")
            try:
                _gmail_api_send(refreshed["access_token"], raw)
            except Exception as e2:
                raise HTTPException(502, f"שליחת המייל נכשלה: {e2}")
        else:
            raise HTTPException(502, f"שליחת המייל נכשלה: {e}")

    # Persist whatever was used (so "set once" sticks), keeping edits
    _save_report_email(u, {
        "recipients": recipients,
        "subject": subj_tmpl if subj_tmpl is not None else cfg["subject"],
        "body":    body_tmpl if body_tmpl is not None else cfg["body"],
    })
    return {"ok": True, "to": recipients, "month": report_month}


# ── Holidays ──────────────────────────────────────────────────────────────────

def _holidays_path(year: int) -> Path:
    return BASE_DIR / "data" / "shared" / f"holidays_{year}.json"

def _load_holidays_cached(year: int):
    p = _holidays_path(year)
    return _rj(p, None) if p.exists() else None

def _save_holidays_cached(year: int, holidays: list):
    p = _holidays_path(year)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(holidays, ensure_ascii=False, indent=2), encoding="utf-8")

def _fetch_and_build_holidays(year: int) -> list:
    import urllib.request
    from datetime import date as _date, timedelta
    url = (f"https://www.hebcal.com/hebcal?v=1&cfg=json&year={year}"
           f"&maj=on&min=on&mod=on&nx=off&ss=off&mf=off&c=off&geo=none&M=on&s=on&i=on")
    try:
        with urllib.request.urlopen(url, timeout=10) as resp:
            items = json.loads(resp.read()).get("items", [])
    except Exception:
        items = []

    yomtov_dates: set = set()
    national_off: set = set()
    national_half: set = set()
    items_by_date: dict = {}
    for item in items:
        d = (item.get("date") or "")[:10]
        if not d:
            continue
        items_by_date.setdefault(d, []).append(item)
        title = item.get("title", "")
        if item.get("yomtov"):
            yomtov_dates.add(d)
        if any(k in title for k in ["Yom HaAtzma'ut", "Yom HaAtzmaut", "Yom ha-Atzma'ut"]):
            national_off.add(d)
        if "Yom HaZikaron" in title:
            national_half.add(d)

    holiday_map: dict = {}

    for d in yomtov_dates:
        item = next((i for i in items_by_date.get(d, []) if i.get("yomtov")), None)
        name = (item or {}).get("hebrew") or (item or {}).get("title") or "חג"
        holiday_map[d] = {"name": name, "type": "off"}

    # Eve of each Yom Tov = half day
    for d in sorted(yomtov_dates):
        eve = (_date.fromisoformat(d) - timedelta(days=1)).isoformat()
        if eve in holiday_map:
            continue
        if _date.fromisoformat(eve).weekday() == 5:  # Saturday
            continue
        yomtov_name = holiday_map.get(d, {}).get("name", "חג")
        holiday_map[eve] = {"name": f"ערב {yomtov_name}", "type": "half"}

    for d in national_off:
        item = next(iter(items_by_date.get(d, [])), None)
        name = (item or {}).get("hebrew") or "יום העצמאות"
        holiday_map[d] = {"name": name, "type": "off"}

    for d in national_half:
        item = next(iter(items_by_date.get(d, [])), None)
        name = (item or {}).get("hebrew") or "יום הזיכרון"
        holiday_map.setdefault(d, {"name": name, "type": "half"})

    return sorted(
        [{"date": d, **v} for d, v in holiday_map.items()],
        key=lambda x: x["date"],
    )

@app.get("/api/holidays/{year}")
async def get_holidays(year: int, refresh: bool = False, _: str = Depends(auth)):
    cached = _load_holidays_cached(year)
    if cached is None or refresh:
        cached = _fetch_and_build_holidays(year)
        _save_holidays_cached(year, cached)
    return cached

@app.put("/api/holidays/{year}")
async def update_holidays(year: int, request: Request, _: str = Depends(auth)):
    body = await request.json()
    _save_holidays_cached(year, body)
    return {"ok": True}

# ── Report months ──────────────────────────────────────────────────────────────

def _date_to_ym(d: str) -> str:
    """DD/MM/YYYY or YYYY-MM-DD → YYYY-MM, empty string if unparseable."""
    d = (d or "").strip()
    if len(d) >= 10 and d[2] == "/" and d[5] == "/":
        return f"{d[6:10]}-{d[3:5]}"
    if len(d) >= 7 and d[4] == "-":
        return d[:7]
    return ""

@app.get("/api/report/months")
async def get_report_months(u: str = Depends(auth)):
    months: set = set()
    for s in _load_saved_shifts(u):
        ym = _date_to_ym(str(s.get("date", "")))
        if ym:
            months.add(ym)
    for s in _load_sales(u):
        ym = _date_to_ym(str(s.get("date", "")))
        if ym:
            months.add(ym)
    return sorted(months, reverse=True)


# ── Salary models + salary report ───────────────────────────────────────────────

# A salary model holds only *rates*; per-worker targets live on the worker.
SALARY_MODEL_FIELDS = [
    "base_hourly",           # שכר לשעה (base 100% rate)
    "bonus_per_sale",        # בונוס מכירות (per approved sale)
    "bonus_chk_1500",        # בונוס אשראי מתגלגל עד 1500
    "bonus_chk_2500",        # בונוס אשראי מתגלגל 1501-2500
    "bonus_chk_4000",        # בונוס אשראי מתגלגל 2501-4000
    "bonus_standing_order",  # בונוס הוראות קבע
    "bonus_arnak_below",     # per ארנוק, when arnakot target NOT reached
    "bonus_arnak_above",     # per ארנוק, when arnakot target reached
    "bonus_extra_per_hour",  # per hour worked, only if sales target reached
    "travel_per_shift",      # נסיעות, per working day
]

def _salary_models_path(u: str) -> Path:
    return _udir(u) / "salary_models.json"

def _load_salary_models(u: str) -> list:
    return _rj(_salary_models_path(u), [])

def _save_salary_models(u: str, models: list):
    _salary_models_path(u).write_text(json.dumps(models, ensure_ascii=False, indent=2), encoding="utf-8")

def _norm_model(body: dict) -> dict:
    def num(v):
        try:
            return float(v)
        except (TypeError, ValueError):
            return 0.0
    out = {"name": (body.get("name") or "").strip()}
    for f in SALARY_MODEL_FIELDS:
        out[f] = num(body.get(f))
    return out

@app.get("/api/salary/models")
async def get_salary_models(u: str = Depends(auth)):
    return _load_salary_models(u)

@app.post("/api/salary/models")
async def add_salary_model(body: dict, u: str = Depends(auth)):
    import uuid as _uuid
    models = _load_salary_models(u)
    m = _norm_model(body)
    if not m["name"]:
        raise HTTPException(400, "שם מודל הוא שדה חובה")
    m["id"] = str(_uuid.uuid4())
    models.append(m)
    _save_salary_models(u, models)
    return {"ok": True, "model": m}

@app.put("/api/salary/models/{model_id}")
async def update_salary_model(model_id: str, body: dict, u: str = Depends(auth)):
    models = _load_salary_models(u)
    for i, m in enumerate(models):
        if m.get("id") == model_id:
            nm = _norm_model(body)
            nm["id"] = model_id
            if not nm["name"]:
                raise HTTPException(400, "שם מודל הוא שדה חובה")
            models[i] = nm
            _save_salary_models(u, models)
            return {"ok": True}
    raise HTTPException(404, "מודל לא נמצא")

@app.delete("/api/salary/models/{model_id}")
async def delete_salary_model(model_id: str, u: str = Depends(auth)):
    models = [m for m in _load_salary_models(u) if m.get("id") != model_id]
    _save_salary_models(u, models)
    return {"ok": True}


def _salary_settings_path(u: str) -> Path:
    return _udir(u) / "salary_settings.json"

_DEFAULT_SALARY_SETTINGS = {"closeness": 2, "hours_125": 9, "hours_150": 11}

def _load_salary_settings(u: str) -> dict:
    return {**_DEFAULT_SALARY_SETTINGS, **_rj(_salary_settings_path(u), {})}

def _save_salary_settings(u: str, s: dict):
    def num(v, d):
        try:
            return float(v)
        except (TypeError, ValueError):
            return d
    clean = {
        "closeness": num(s.get("closeness"), _DEFAULT_SALARY_SETTINGS["closeness"]),
        "hours_125": num(s.get("hours_125"), _DEFAULT_SALARY_SETTINGS["hours_125"]),
        "hours_150": num(s.get("hours_150"), _DEFAULT_SALARY_SETTINGS["hours_150"]),
    }
    _salary_settings_path(u).write_text(json.dumps(clean, ensure_ascii=False, indent=2), encoding="utf-8")
    return clean


def _salary_runs_path(u: str) -> Path:
    return _udir(u) / "salary_runs.json"

def _load_salary_runs(u: str) -> dict:
    return _rj(_salary_runs_path(u), {})

def _save_salary_run(u: str, month: str, run: dict):
    runs = _load_salary_runs(u)
    runs[month] = run
    _salary_runs_path(u).write_text(json.dumps(runs, ensure_ascii=False, indent=2), encoding="utf-8")


def _salary_aggregate(u: str, month: str) -> dict:
    """Per-worker monthly aggregates keyed by worker id.
    Returns {worker_id: {...counts...}} plus the worker record."""
    workers   = _load_workers(u)
    shifts    = _load_saved_shifts(u)
    sales     = _load_sales(u)
    arnakot   = _load_arnakot(u)
    resolve   = _worker_resolver(u)

    # Managers are excluded from salary entirely (their pay is handled separately)
    by_name = {
        w.get("full_name", "").strip(): w
        for w in workers
        if w.get("full_name") and w.get("rank") != "manager"
    }

    def in_month(d):
        return _date_to_ym(str(d)) == month

    agg: dict = {}
    def slot(w):
        wid = w["id"]
        if wid not in agg:
            agg[wid] = {
                "worker": w,
                "day_hours": {},   # date -> summed hours
                "approved": 0, "rev1500": 0, "rev2500": 0, "rev4000": 0,
                "so": 0, "arnakot": 0,
            }
        return agg[wid]

    # Hours (by resolved worker, summed per day)
    for s in shifts:
        if not in_month(s.get("date", "")):
            continue
        key = (s.get("worker_key") or "").strip() or resolve(s.get("worker_name", ""))
        w = by_name.get(key)
        if not w:
            continue
        d = str(s.get("date", ""))
        slot(w)["day_hours"][d] = slot(w)["day_hours"].get(d, 0.0) + float(s.get("hours") or 0)

    # Sales
    for s in sales:
        if not in_month(s.get("date", "")):
            continue
        name = ((s.get("first_name") or "") + " " + (s.get("last_name") or "")).strip()
        w = by_name.get(resolve(name))
        if not w:
            continue
        st = slot(w)
        if s.get("approved"):       st["approved"] += 1
        if s.get("revolving_1500"): st["rev1500"]  += 1
        if s.get("revolving_2500"): st["rev2500"]  += 1
        if s.get("revolving_4000"): st["rev4000"]  += 1
        if s.get("standing_order"): st["so"]       += 1

    # Arnakot
    for a in arnakot:
        if not in_month(a.get("date", "")):
            continue
        w = by_name.get(resolve(a.get("name", "")))
        if not w:
            continue
        slot(w)["arnakot"] += 1

    return agg


def _num(v, d=0.0):
    try:
        return float(v)
    except (TypeError, ValueError):
        return d


def _compute_salary(u: str, month: str, settings: dict, overrides: dict, manager_bonuses: dict):
    """Returns list of per-worker salary rows (one per worker that has a model)."""
    models_by_id = {m["id"]: m for m in _load_salary_models(u)}
    agg = _salary_aggregate(u, month)

    h125_thr = _num(settings.get("hours_125"), _DEFAULT_SALARY_SETTINGS["hours_125"])
    h150_thr = _num(settings.get("hours_150"), _DEFAULT_SALARY_SETTINGS["hours_150"])
    if h150_thr < h125_thr:
        h150_thr = h125_thr

    rows = []
    for wid, a in agg.items():
        w = a["worker"]
        model = models_by_id.get(w.get("salary_model"))
        if not model:
            continue  # only workers with an assigned model get a salary row

        sales_target   = _num(w.get("sales_target"))
        arnakot_target = _num(w.get("arnakot_target"))
        approved       = a["approved"]
        arnakot_cnt    = a["arnakot"]

        auto_sales   = sales_target > 0 and approved >= sales_target
        auto_arnak   = arnakot_target > 0 and arnakot_cnt >= arnakot_target
        ov           = overrides.get(wid, {})
        sales_reached   = ov["sales"]   if isinstance(ov.get("sales"), bool)   else auto_sales
        arnakot_reached = ov["arnakot"] if isinstance(ov.get("arnakot"), bool) else auto_arnak

        # Daily overtime tiering
        h100 = h125 = h150 = 0.0
        for d, dh in a["day_hours"].items():
            h100 += min(dh, h125_thr)
            h125 += max(0.0, min(dh, h150_thr) - h125_thr)
            h150 += max(0.0, dh - h150_thr)
        total_hours = h100 + h125 + h150
        working_days = len(a["day_hours"])

        base = _num(model.get("base_hourly"))
        mgr_bonus = _num(manager_bonuses.get(wid))

        d_travel = _num(model.get("travel_per_shift")) * working_days
        e_sales  = _num(model.get("bonus_per_sale")) * approved
        f_chk1   = _num(model.get("bonus_chk_1500")) * a["rev1500"]
        h_chk2   = _num(model.get("bonus_chk_2500")) * a["rev2500"]
        i_chk3   = _num(model.get("bonus_chk_4000")) * a["rev4000"]
        j_so     = _num(model.get("bonus_standing_order")) * a["so"]
        arnak_rate = _num(model.get("bonus_arnak_above")) if arnakot_reached else _num(model.get("bonus_arnak_below"))
        k_arnak  = arnak_rate * arnakot_cnt
        l_extra  = (_num(model.get("bonus_extra_per_hour")) * total_hours) if sales_reached else 0.0

        pay_hours = base * h100 + base * 1.25 * h125 + base * 1.5 * h150
        total = pay_hours + d_travel + e_sales + f_chk1 + mgr_bonus + h_chk2 + i_chk3 + j_so + k_arnak + l_extra

        rows.append({
            "worker_id": wid,
            "id_number": w.get("id_number", ""),
            "name": w.get("full_name", ""),
            "model_name": model.get("name", ""),
            "base_hourly": round(base, 2),
            "travel": round(d_travel, 2),
            "bonus_sales": round(e_sales, 2),
            "bonus_chk1500": round(f_chk1, 2),
            "manager_bonus": round(mgr_bonus, 2),
            "bonus_chk2500": round(h_chk2, 2),
            "bonus_chk4000": round(i_chk3, 2),
            "bonus_standing_order": round(j_so, 2),
            "bonus_arnakot": round(k_arnak, 2),
            "bonus_extra": round(l_extra, 2),
            "working_days": working_days,
            "hours_100": round(h100, 2),
            "hours_125": round(h125, 2),
            "hours_150": round(h150, 2),
            "total": round(total, 2),
            # extra context for the UI
            "approved_sales": approved, "sales_target": sales_target,
            "arnakot_count": arnakot_cnt, "arnakot_target": arnakot_target,
            "sales_reached": sales_reached, "arnakot_reached": arnakot_reached,
        })

    rows.sort(key=lambda r: r["name"])
    return rows


@app.get("/api/salary/prepare")
async def salary_prepare(u: str = Depends(auth), month: str = Query(...)):
    """Per-worker aggregates + auto reached flags + any saved decisions,
    so the frontend can drive the closeness review before computing."""
    models = _load_salary_models(u)
    models_by_id = {m["id"]: m for m in models}
    agg = _salary_aggregate(u, month)

    workers_out = []
    for wid, a in agg.items():
        w = a["worker"]
        model = models_by_id.get(w.get("salary_model"))
        sales_target   = _num(w.get("sales_target"))
        arnakot_target = _num(w.get("arnakot_target"))
        workers_out.append({
            "worker_id": wid,
            "id_number": w.get("id_number", ""),
            "name": w.get("full_name", ""),
            "has_model": bool(model),
            "model_name": model.get("name", "") if model else "",
            "sales_target": sales_target,
            "arnakot_target": arnakot_target,
            "approved_sales": a["approved"],
            "arnakot_count": a["arnakot"],
            "auto_sales_reached": sales_target > 0 and a["approved"] >= sales_target,
            "auto_arnakot_reached": arnakot_target > 0 and a["arnakot"] >= arnakot_target,
        })
    workers_out.sort(key=lambda x: x["name"])

    saved = _load_salary_runs(u).get(month)
    return {
        "month": month,
        "models_exist": bool(models),
        "settings": _load_salary_settings(u),
        "workers": workers_out,
        "saved": saved,
    }


@app.post("/api/salary/compute")
async def salary_compute(body: dict, u: str = Depends(auth)):
    from datetime import datetime
    month = body.get("month")
    if not month:
        raise HTTPException(400, "חסר חודש")
    settings = _save_salary_settings(u, body.get("settings") or {})
    overrides = body.get("overrides") or {}
    manager_bonuses = body.get("manager_bonuses") or {}
    rows = _compute_salary(u, month, settings, overrides, manager_bonuses)
    run = {
        "month": month,
        "settings": settings,
        "overrides": overrides,
        "manager_bonuses": manager_bonuses,
        "rows": rows,
        "computed_at": datetime.now().isoformat(timespec="seconds"),
    }
    if body.get("save", True):
        _save_salary_run(u, month, run)
    return run


@app.get("/api/salary/export")
async def salary_export(u: str = Depends(auth), month: str = Query(...)):
    run = _load_salary_runs(u).get(month)
    if not run:
        raise HTTPException(404, "אין חישוב שמור לחודש זה. יש לחשב שכר תחילה.")
    rows = run.get("rows", [])

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook()
    ws = wb.active
    ws.title = "דוח שכר"
    ws.sheet_view.rightToLeft = True

    headers = [
        "ת.ז", "שם עובד", "שכר לשעה", "נסיעות", "בונוס מכירות",
        "בונוס אשראי מתגלגל עד 1500", "בונוס מנהל", "בונוס אשראי מתגלגל 1501-2500",
        "בונוס אשראי מתגלגל 2501-4000", "בונוס הוראות קבע", "בונוס ארנוקים",
        "בונוס אקסטרה", "ימי עבודה", "שעות 100%", "שעות 125%", "שעות 150%",
        "סה\"כ שכר",
    ]
    keys = [
        "id_number", "name", "base_hourly", "travel", "bonus_sales",
        "bonus_chk1500", "manager_bonus", "bonus_chk2500",
        "bonus_chk4000", "bonus_standing_order", "bonus_arnakot",
        "bonus_extra", "working_days", "hours_100", "hours_125", "hours_150",
        "total",
    ]
    hdr_fill = PatternFill("solid", fgColor="2F5496")
    hdr_font = Font(bold=True, color="FFFFFF", size=10)
    center   = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin     = Side(style="thin", color="B0B0B0")
    border   = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.append(headers)
    for c in ws[1]:
        c.font = hdr_font; c.fill = hdr_fill; c.alignment = center; c.border = border
    total_col = len(keys)
    for r in rows:
        ws.append([r.get(k, "") for k in keys])
        for ci, c in enumerate(ws[ws.max_row], 1):
            c.alignment = center; c.border = border
            if ci == total_col:
                c.font = Font(bold=True)
                c.fill = PatternFill("solid", fgColor="FFF2CC")
    for col in ws.columns:
        w = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(w + 3, 22)

    out = Path(tempfile.mktemp(suffix=".xlsx"))
    wb.save(str(out))
    data = out.read_bytes()
    out.unlink(missing_ok=True)
    fname = f"salary_report_{month}.xlsx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ── נתיב האור (Netiv HaOr) — guide work arrangements ──────────────────────────
# Data is shared across the managers who have access, not per-user.

NETIV_TASKS_FILE   = BASE_DIR / "data" / "shared" / "netiv_tasks.json"
NETIV_WORKERS_FILE = BASE_DIR / "data" / "shared" / "netiv_workers.json"

# Excel header text -> internal field. Matched case-insensitively by substring.
NETIV_COLUMNS = [
    ("guide_name",    ["שם מדריך"]),
    ("guide_id",      ["תז מדריך (לסופרפורם)", "תז מדריך", "ת.ז מדריך"]),
    ("guide_phone",   ["טלפון מדריך"]),
    ("school",        ["שם בית ספר", "בית ספר"]),
    ("address",       ["כתובת"]),
    ("city",          ["עיר"]),
    ("contact_name",  ["שם איש קשר"]),
    ("contact_phone", ["טלפון איש קשר"]),
    ("contact_email", ["מייל איש קשר"]),
    ("region",        ["אזור"]),
    ("institution",   ["סמל מוסד"]),
    ("date",          ["תאריך"]),
    ("class_name",    ["כיתה"]),
    ("hour",          ["שעות", "שעה"]),
    ("lesson_num",    ["מס' הדרכה", "מס הדרכה"]),
    ("notes",         ["הערות"]),
    ("done",          ["בוצע שיעור"]),
]

SESSION_MINUTES = 45  # every session is a fixed 45 minutes


def _load_netiv_tasks() -> list:
    return _rj(NETIV_TASKS_FILE, [])

def _save_netiv_tasks(tasks: list):
    _wj(NETIV_TASKS_FILE, tasks)

def _load_netiv_workers() -> list:
    return _rj(NETIV_WORKERS_FILE, [])

def _save_netiv_workers(workers: list):
    _wj(NETIV_WORKERS_FILE, workers)


def _netiv_norm_time(val) -> str:
    """Normalize a session start time to 'HH:MM' (24h). Accepts '12:00 PM',
    '13:00', datetime/time objects and Excel fractions."""
    if val is None or val == "":
        return ""
    if hasattr(val, "strftime") and hasattr(val, "hour"):
        return val.strftime("%H:%M")
    if isinstance(val, (int, float)) and 0 <= float(val) <= 1:
        total = round(float(val) * 86400)
        return f"{(total // 3600) % 24:02d}:{(total % 3600) // 60:02d}"
    s = str(val).strip().upper().replace(".", "")
    m = re.match(r"^(\d{1,2}):(\d{2})(?::\d{2})?\s*(AM|PM)?$", s)
    if not m:
        return str(val).strip()
    hh, mm, ampm = int(m.group(1)), int(m.group(2)), m.group(3)
    if ampm == "PM" and hh != 12:
        hh += 12
    elif ampm == "AM" and hh == 12:
        hh = 0
    return f"{hh:02d}:{mm:02d}"


def _netiv_norm_date(val) -> str:
    """Normalize a date to DD/MM/YYYY."""
    if val is None or val == "":
        return ""
    if hasattr(val, "strftime"):
        return val.strftime("%d/%m/%Y")
    s = str(val).strip()
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        return f"{m.group(3)}/{m.group(2)}/{m.group(1)}"
    return s


def _netiv_week_start(ddmmyyyy: str) -> str:
    """Sunday of that date's week, as YYYY-MM-DD (Israeli week: Sun-Sat)."""
    from datetime import date as _date, timedelta as _td
    try:
        d, m, y = ddmmyyyy.split("/")
        dt = _date(int(y), int(m), int(d))
    except Exception:
        return ""
    # Python weekday(): Mon=0 .. Sun=6  ->  days since Sunday
    return (dt - _td(days=(dt.weekday() + 1) % 7)).isoformat()


def _netiv_week_label(week_start_iso: str) -> str:
    from datetime import date as _date, timedelta as _td
    try:
        y, m, d = map(int, week_start_iso.split("-"))
        start = _date(y, m, d)
    except Exception:
        return week_start_iso
    end = start + _td(days=6)
    return f"{start.strftime('%d/%m')}–{end.strftime('%d/%m/%Y')}"


@app.post("/api/netiv/upload")
async def netiv_upload(u: str = Depends(auth), file: UploadFile = File(...)):
    import pandas as pd, uuid as _uuid
    contents = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(contents)
        tmp_path = Path(tmp.name)
    try:
        raw = pd.read_excel(str(tmp_path), header=None)

        # The monday.com export has preamble rows; find the real header row
        header_idx = None
        for i in range(min(len(raw), 30)):
            row_text = " ".join(str(v) for v in raw.iloc[i].tolist() if v is not None)
            if "שם מדריך" in row_text:
                header_idx = i
                break
        if header_idx is None:
            raise HTTPException(400, "לא נמצאה שורת כותרת עם 'שם מדריך' בקובץ")

        headers = [str(v).strip() if v is not None else "" for v in raw.iloc[header_idx].tolist()]
        col_of: dict = {}
        for field, keys in NETIV_COLUMNS:
            for ci, h in enumerate(headers):
                if ci in col_of.values() or not h or h == "nan":
                    continue
                if any(k.lower() in h.lower() for k in keys):
                    col_of[field] = ci
                    break

        def cell(row, field) -> str:
            ci = col_of.get(field)
            if ci is None or ci >= len(row):
                return ""
            v = row.iloc[ci]
            if v is None:
                return ""
            try:
                if pd.isna(v):
                    return ""
            except (TypeError, ValueError):
                pass
            return v if field in ("date", "hour") else str(v).strip()

        tasks = []
        for i in range(header_idx + 1, len(raw)):
            row = raw.iloc[i]
            guide  = str(cell(row, "guide_name") or "").strip()
            school = str(cell(row, "school") or "").strip()
            date_s = _netiv_norm_date(cell(row, "date"))
            # Skip monday.com summary/aggregation rows: they carry a date but
            # no guide and no school
            if not date_s or (not guide and not school):
                continue
            tasks.append({
                "id":            str(_uuid.uuid4()),
                "guide_name":    guide,
                "guide_id":      str(cell(row, "guide_id") or "").strip(),
                "guide_phone":   str(cell(row, "guide_phone") or "").strip(),
                "school":        school,
                "address":       str(cell(row, "address") or "").strip(),
                "city":          str(cell(row, "city") or "").strip(),
                "contact_name":  str(cell(row, "contact_name") or "").strip(),
                "contact_phone": str(cell(row, "contact_phone") or "").strip(),
                "contact_email": str(cell(row, "contact_email") or "").strip(),
                "region":        str(cell(row, "region") or "").strip(),
                "institution":   str(cell(row, "institution") or "").strip(),
                "date":          date_s,
                "week":          _netiv_week_start(date_s),
                "class_name":    str(cell(row, "class_name") or "").strip(),
                "hour":          _netiv_norm_time(cell(row, "hour")),
                "lesson_num":    str(cell(row, "lesson_num") or "").strip(),
                "notes":         str(cell(row, "notes") or "").strip(),
                "done":          str(cell(row, "done") or "").strip(),
            })

        _save_netiv_tasks(tasks)  # a new upload replaces the previous file
        added = _netiv_sync_workers(tasks)
        weeks = sorted({t["week"] for t in tasks if t["week"]})
        return {"ok": True, "count": len(tasks), "workers_added": added, "weeks": len(weeks)}
    finally:
        tmp_path.unlink(missing_ok=True)


def _netiv_sync_workers(tasks: list) -> int:
    """Auto-add guides seen in the tasks to the workers list. Existing workers
    keep their (possibly hand-corrected) phone number."""
    import uuid as _uuid
    workers = _load_netiv_workers()
    matcher = NameMatcher([w.get("full_name", "") for w in workers])
    by_name = {normalize_match_text(w.get("full_name", "")): w for w in workers}
    added = 0
    for t in tasks:
        name = normalize_match_text(t.get("guide_name", ""))
        if not name:
            continue
        if matcher.resolve(name) or name in by_name:
            continue
        w = {
            "id":        str(_uuid.uuid4()),
            "full_name": t["guide_name"].strip(),
            "id_number": t.get("guide_id", ""),
            "phone":     t.get("guide_phone", ""),
            "notes":     "",
        }
        workers.append(w)
        by_name[name] = w
        matcher = NameMatcher([x.get("full_name", "") for x in workers])
        added += 1
    if added:
        _save_netiv_workers(workers)
    return added


UNASSIGNED_GUIDE = "ללא מדריך"


def _netiv_build_week(week: str = None):
    """Arrangement for one week: per-guide blocks plus warnings."""
    tasks = _load_netiv_tasks()
    weeks = sorted({t["week"] for t in tasks if t.get("week")})
    if not weeks:
        return {"weeks": [], "week": "", "week_label": "", "hours": [], "guides": [], "warnings": []}
    sel = week if week in weeks else weeks[-1]
    in_week = [t for t in tasks if t.get("week") == sel]

    workers = _load_netiv_workers()
    phone_by_name = {normalize_match_text(w.get("full_name", "")): w.get("phone", "") for w in workers}
    id_by_name    = {normalize_match_text(w.get("full_name", "")): w.get("id_number", "") for w in workers}

    hours = sorted({t["hour"] for t in in_week if t.get("hour")})

    # guide -> date -> hour -> [tasks]
    guides: dict = {}
    for t in in_week:
        gname = t.get("guide_name", "").strip() or UNASSIGNED_GUIDE
        g = guides.setdefault(gname, {})
        g.setdefault(t.get("date", ""), {}).setdefault(t.get("hour", ""), []).append(t)

    warnings = []
    for gname, dates in guides.items():
        for d, by_hour in dates.items():
            for h, items in by_hour.items():
                if len(items) > 1:
                    warnings.append({
                        "type": "overlap",
                        "guide": gname, "date": d, "hour": h,
                        "text": f"{gname} — {d} בשעה {h}: {len(items)} הדרכות באותו זמן "
                                f"({', '.join(i.get('school','') or '—' for i in items)})",
                    })
    for t in in_week:
        if not t.get("guide_name", "").strip():
            warnings.append({
                "type": "no_guide", "guide": UNASSIGNED_GUIDE,
                "date": t.get("date", ""), "hour": t.get("hour", ""),
                "text": f"הדרכה ללא מדריך — {t.get('date','')} {t.get('hour','')} "
                        f"{t.get('school','') or ''} ({t.get('city','') or ''})",
            })
        if str(t.get("done", "")).strip().upper() == "V":
            warnings.append({
                "type": "already_done", "guide": t.get("guide_name", "") or UNASSIGNED_GUIDE,
                "date": t.get("date", ""), "hour": t.get("hour", ""),
                "text": f"הדרכה כבר מסומנת כבוצעה — {t.get('guide_name','') or UNASSIGNED_GUIDE}, "
                        f"{t.get('date','')} {t.get('hour','')} {t.get('school','') or ''}",
            })

    def date_key(d):
        try:
            dd, mm, yy = d.split("/")
            return f"{yy}-{mm}-{dd}"
        except Exception:
            return d

    out_guides = []
    for gname in sorted(guides, key=lambda n: (n == UNASSIGNED_GUIDE, n)):
        key = normalize_match_text(gname)
        rows = []
        for d in sorted(guides[gname], key=date_key):
            cells = []
            for h in hours:
                cells.append([
                    {
                        "school": t.get("school", ""), "city": t.get("city", ""),
                        "address": t.get("address", ""), "contact_name": t.get("contact_name", ""),
                        "contact_phone": t.get("contact_phone", ""), "class_name": t.get("class_name", ""),
                        "notes": t.get("notes", ""),
                    }
                    for t in guides[gname][d].get(h, [])
                ])
            rows.append({"date": d, "cells": cells})
        out_guides.append({
            "name": gname,
            "guide_id": id_by_name.get(key, "") or next((t.get("guide_id", "") for d in guides[gname]
                        for h in guides[gname][d] for t in guides[gname][d][h] if t.get("guide_id")), ""),
            "phone": phone_by_name.get(key, ""),
            "rows": rows,
        })

    return {
        "weeks": [{"value": w, "label": _netiv_week_label(w)} for w in weeks],
        "week": sel, "week_label": _netiv_week_label(sel),
        "hours": hours, "guides": out_guides, "warnings": warnings,
        "session_minutes": SESSION_MINUTES,
    }


@app.get("/api/netiv/arrangement")
async def netiv_arrangement(_: str = Depends(auth), week: str = Query(None)):
    return _netiv_build_week(week)


@app.get("/api/netiv/workers")
async def netiv_get_workers(_: str = Depends(auth)):
    return _load_netiv_workers()


@app.post("/api/netiv/workers")
async def netiv_add_worker(body: dict, _: str = Depends(auth)):
    import uuid as _uuid
    workers = _load_netiv_workers()
    body["id"] = str(_uuid.uuid4())
    if not (body.get("full_name") or "").strip():
        raise HTTPException(400, "שם מלא הוא שדה חובה")
    workers.append(body)
    _save_netiv_workers(workers)
    return {"ok": True}


@app.put("/api/netiv/workers/{worker_id}")
async def netiv_update_worker(worker_id: str, body: dict, _: str = Depends(auth)):
    workers = _load_netiv_workers()
    for i, w in enumerate(workers):
        if w.get("id") == worker_id:
            body["id"] = worker_id
            workers[i] = body
            _save_netiv_workers(workers)
            return {"ok": True}
    raise HTTPException(404, "עובד לא נמצא")


@app.delete("/api/netiv/workers/{worker_id}")
async def netiv_delete_worker(worker_id: str, _: str = Depends(auth)):
    _save_netiv_workers([w for w in _load_netiv_workers() if w.get("id") != worker_id])
    return {"ok": True}


@app.get("/api/netiv/arrangement/pdf")
async def netiv_arrangement_pdf(_: str = Depends(auth), week: str = Query(None),
                                guide: str = Query(None)):
    """Weekly arrangement as PDF. With ?guide=<name> it contains only that
    guide — which is what each worker will receive."""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from bidi.algorithm import get_display
    import io

    data = _netiv_build_week(week)
    guides = data["guides"]
    if guide:
        guides = [g for g in guides if g["name"] == guide]
        if not guides:
            raise HTTPException(404, "לא נמצא סידור עבודה למדריך זה בשבוע הנבחר")

    if "DejaVu" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVu", str(BASE_DIR / "fonts" / "DejaVuSans.ttf")))
        pdfmetrics.registerFont(TTFont("DejaVu-Bold", str(BASE_DIR / "fonts" / "DejaVuSans-Bold.ttf")))

    def heb(t):
        return get_display(str(t))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            rightMargin=8*mm, leftMargin=8*mm, topMargin=10*mm, bottomMargin=10*mm)
    title_style = ParagraphStyle("t", fontName="DejaVu-Bold", fontSize=13, alignment=2)
    sub_style   = ParagraphStyle("s", fontName="DejaVu", fontSize=9, alignment=2)
    cell_style  = ParagraphStyle("c", fontName="DejaVu", fontSize=6.5, alignment=2, leading=8)
    elements = []

    hours = data["hours"]
    for gi, g in enumerate(guides):
        if gi:
            elements.append(Spacer(1, 8*mm))
        header_bits = [g["name"]]
        if g.get("guide_id"):
            header_bits.append(f"ת.ז: {g['guide_id']}")
        elements.append(Paragraph(heb(" | ".join(header_bits)), title_style))
        elements.append(Paragraph(heb(f"סידור עבודה — שבוע {data['week_label']} (כל הדרכה {SESSION_MINUTES} דקות)"), sub_style))
        elements.append(Spacer(1, 3*mm))

        # RTL: reverse column order so "תאריך" ends up rightmost
        head = ["תאריך"] + hours
        table_data = [[Paragraph(heb(h), cell_style) for h in reversed(head)]]
        for row in g["rows"]:
            line = [row["date"]]
            for cell in row["cells"]:
                line.append("\n———\n".join(_netiv_cell_text(item) for item in cell) if cell else "")
            table_data.append([Paragraph(heb(c).replace("\n", "<br/>"), cell_style) for c in reversed(line)])

        col_w = [26*mm] + [(255 - 26) / max(len(hours), 1) * mm] * len(hours)
        t = Table(table_data, colWidths=list(reversed(col_w)), repeatRows=1)
        t.setStyle(TableStyle([
            ("FONTNAME",   (0, 0), (-1, -1), "DejaVu"),
            ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#2F5496")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), rl_colors.white),
            ("GRID",       (0, 0), (-1, -1), 0.4, rl_colors.grey),
            ("VALIGN",     (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(t)

    doc.build(elements)
    pdf = buf.getvalue()
    fname = f"netiv_{(guide or 'all').replace(' ', '_')}_{data['week']}.pdf"
    return Response(content=pdf, media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


def _netiv_cell_text(item: dict) -> str:
    """One session rendered as labelled lines (shared by PDF and Excel)."""
    parts = [
        ("עיר", item.get("city")),
        ("בית ספר", item.get("school")),
        ("כתובת", item.get("address")),
        ("איש קשר", item.get("contact_name")),
        ("טלפון", item.get("contact_phone")),
        ("כיתה", item.get("class_name")),
        ("הערות", item.get("notes")),
    ]
    return "\n".join(f"{label}: {val}" for label, val in parts if val)


# ── Recruiter Analysis ────────────────────────────────────────────────────────

RECRUITER_CONFIG_FILE = BASE_DIR / "data" / "shared" / "recruiter_config.json"
RECRUITER_DATA_FILE = BASE_DIR / "data" / "shared" / "recruiter_data.json"

_DEFAULT_RECRUITER_CONFIG: dict = {
    "recruiters": [],
    "long_call_threshold_minutes": 8,
    "repeat_call_threshold": 2,
    "default_days_back": 7,
}

def _load_recruiter_config() -> dict:
    return _rj(RECRUITER_CONFIG_FILE, dict(_DEFAULT_RECRUITER_CONFIG))

def _save_recruiter_config(cfg: dict):
    _wj(RECRUITER_CONFIG_FILE, cfg)

# ── WhatsApp Webhook ──────────────────────────────────────────────────────────
WHATSAPP_VERIFY_TOKEN = os.environ.get("WHATSAPP_VERIFY_TOKEN", "electra_target_verify_2024")
_whatsapp_messages: list = []  # stores last received shift messages

@app.get("/api/whatsapp/webhook")
async def whatsapp_verify(request: Request):
    """Meta webhook verification handshake."""
    params = dict(request.query_params)
    if params.get("hub.verify_token") == WHATSAPP_VERIFY_TOKEN and params.get("hub.mode") == "subscribe":
        return Response(content=params.get("hub.challenge", ""), media_type="text/plain")
    raise HTTPException(403, "Invalid verify token")

@app.post("/api/whatsapp/webhook")
async def whatsapp_receive(request: Request):
    """Receive incoming WhatsApp messages."""
    data = await request.json()
    try:
        for entry in data.get("entry", []):
            for change in entry.get("changes", []):
                messages = change.get("value", {}).get("messages", [])
                for msg in messages:
                    if msg.get("type") == "text":
                        text = msg["text"]["body"]
                        if text.strip().startswith("!משמרות"):
                            content = text[len("!משמרות"):].strip()
                            _whatsapp_messages.insert(0, {"text": content, "timestamp": msg.get("timestamp","")})
                            _whatsapp_messages[:] = _whatsapp_messages[:10]  # keep last 10
    except Exception:
        pass
    return {"status": "ok"}

@app.get("/api/whatsapp/latest-message")
async def whatsapp_latest(u: str = Depends(auth)):
    if not _whatsapp_messages:
        return {"ok": False, "msg": "לא נמצאה הודעת משמרות"}
    return {"ok": True, "text": _whatsapp_messages[0]["text"]}


# ── Gmail OAuth ───────────────────────────────────────────────────────────────
GMAIL_CLIENT_ID     = os.environ.get("GOOGLE_CLIENT_ID", "")
GMAIL_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")
GMAIL_REDIRECT_URI  = "https://tender-scanner.up.railway.app/auth/gmail/callback"
GMAIL_SCOPES        = "https://www.googleapis.com/auth/gmail.readonly"
GMAIL_SEND_SCOPE    = "https://www.googleapis.com/auth/gmail.send"
# Railway blocks outbound SMTP, so report emails go out via the Gmail HTTPS API
# from one fixed sender account, connected once via OAuth and stored under this key.
REPORT_SENDER_KEY   = "__report_sender__"
_GMAIL_TOKENS_FILE = BASE_DIR / "data" / "shared" / "gmail_tokens.json"

def _load_gmail_tokens() -> dict:
    return _rj(_GMAIL_TOKENS_FILE, {})

def _save_gmail_token(u: str, tokens: dict):
    if "expires_in" in tokens:
        tokens["expires_at"] = time.time() + tokens["expires_in"]
    all_tokens = _load_gmail_tokens()
    all_tokens[u] = tokens
    _GMAIL_TOKENS_FILE.write_text(json.dumps(all_tokens, ensure_ascii=False, indent=2), encoding="utf-8")

def _clear_gmail_token(u: str):
    all_tokens = _load_gmail_tokens()
    if u in all_tokens:
        del all_tokens[u]
        _GMAIL_TOKENS_FILE.write_text(json.dumps(all_tokens, ensure_ascii=False, indent=2), encoding="utf-8")

def _refresh_gmail_tokens(u: str, tokens: dict) -> Optional[dict]:
    """Exchange the stored refresh_token for a fresh access_token.
    Google only returns a refresh_token on the first consent grant, so it's
    carried over here. Returns None (and clears the stored token) if refresh
    is impossible or the token was revoked."""
    import urllib.request, urllib.parse, json as _json
    refresh_token = tokens.get("refresh_token")
    if not refresh_token:
        _clear_gmail_token(u)
        return None
    data = urllib.parse.urlencode({
        "refresh_token": refresh_token,
        "client_id": GMAIL_CLIENT_ID,
        "client_secret": GMAIL_CLIENT_SECRET,
        "grant_type": "refresh_token",
    }).encode()
    req = urllib.request.Request("https://oauth2.googleapis.com/token", data=data,
        headers={"Content-Type": "application/x-www-form-urlencoded"}, method="POST")
    try:
        with urllib.request.urlopen(req) as resp:
            new_tokens = _json.loads(resp.read())
    except Exception:
        _clear_gmail_token(u)
        return None
    new_tokens.setdefault("refresh_token", refresh_token)
    _save_gmail_token(u, new_tokens)
    return new_tokens

def _get_valid_gmail_tokens(u: str) -> Optional[dict]:
    """Returns a token dict with a currently-valid access_token, refreshing
    proactively when the stored one is missing/near expiry. None means the
    user genuinely needs to reconnect."""
    tokens = _load_gmail_tokens().get(u)
    if not tokens:
        return None
    expires_at = tokens.get("expires_at", 0)
    if expires_at - time.time() > 60:
        return tokens
    return _refresh_gmail_tokens(u, tokens)

@app.get("/auth/gmail/connect")
async def gmail_connect(u: str = Depends(auth)):
    import urllib.parse
    params = urllib.parse.urlencode({
        "client_id": GMAIL_CLIENT_ID,
        "redirect_uri": GMAIL_REDIRECT_URI,
        "response_type": "code",
        "scope": GMAIL_SCOPES,
        "access_type": "offline",
        "prompt": "consent",
        "state": u,
    })
    from fastapi.responses import RedirectResponse
    return RedirectResponse(f"https://accounts.google.com/o/oauth2/v2/auth?{params}")

@app.get("/auth/gmail/callback")
async def gmail_callback(code: str, state: str):
    import urllib.request, urllib.parse, json as _json
    data = urllib.parse.urlencode({
        "code": code,
        "client_id": GMAIL_CLIENT_ID,
        "client_secret": GMAIL_CLIENT_SECRET,
        "redirect_uri": GMAIL_REDIRECT_URI,
        "grant_type": "authorization_code",
    }).encode()
    req = urllib.request.Request("https://oauth2.googleapis.com/token", data=data,
        headers={"Content-Type": "application/x-www-form-urlencoded"}, method="POST")
    with urllib.request.urlopen(req) as resp:
        tokens = _json.loads(resp.read())

    # For the report-sender connection, remember which address was connected
    signal = "gmail_connected"
    if state == REPORT_SENDER_KEY:
        signal = "report_gmail_connected"
        try:
            preq = urllib.request.Request(
                "https://gmail.googleapis.com/gmail/v1/users/me/profile",
                headers={"Authorization": f"Bearer {tokens.get('access_token','')}"})
            with urllib.request.urlopen(preq) as pr:
                tokens["email"] = _json.loads(pr.read()).get("emailAddress", "")
        except Exception:
            pass

    _save_gmail_token(state, tokens)
    from fastapi.responses import HTMLResponse
    return HTMLResponse(f"<script>window.close();window.opener&&window.opener.postMessage('{signal}','*')</script>✅ Gmail מחובר! ניתן לסגור חלון זה.")

@app.get("/auth/report-gmail/connect")
async def report_gmail_connect(u: str = Depends(auth)):
    import urllib.parse
    # The site account does BOTH jobs: read the clock2go attendance mail and
    # send the reports — so request read + send together.
    params = urllib.parse.urlencode({
        "client_id": GMAIL_CLIENT_ID,
        "redirect_uri": GMAIL_REDIRECT_URI,
        "response_type": "code",
        "scope": f"{GMAIL_SCOPES} {GMAIL_SEND_SCOPE}",
        "access_type": "offline",
        "prompt": "consent",
        "state": REPORT_SENDER_KEY,
    })
    from fastapi.responses import RedirectResponse
    return RedirectResponse(f"https://accounts.google.com/o/oauth2/v2/auth?{params}")

@app.post("/auth/report-gmail/disconnect")
async def report_gmail_disconnect(u: str = Depends(auth)):
    _clear_gmail_token(REPORT_SENDER_KEY)
    return {"ok": True}

def _report_sender_email() -> str:
    return (_load_gmail_tokens().get(REPORT_SENDER_KEY, {}) or {}).get("email", "")

def _gmail_account_for(u: str) -> Optional[str]:
    """Which stored account to use for a Gmail read op: the shared site account
    if connected, otherwise the per-user connection (legacy fallback)."""
    if _get_valid_gmail_tokens(REPORT_SENDER_KEY):
        return REPORT_SENDER_KEY
    if _get_valid_gmail_tokens(u):
        return u
    return None

def _gmail_api_send(access_token: str, raw_b64: str) -> dict:
    """POST a base64url-encoded RFC822 message to the Gmail send API."""
    import urllib.request, json as _json
    req = urllib.request.Request(
        "https://gmail.googleapis.com/gmail/v1/users/me/messages/send",
        data=_json.dumps({"raw": raw_b64}).encode(),
        headers={"Authorization": f"Bearer {access_token}", "Content-Type": "application/json"},
        method="POST")
    with urllib.request.urlopen(req, timeout=30) as r:
        return _json.loads(r.read())

@app.get("/auth/gmail/status")
async def gmail_status(u: str = Depends(auth)):
    # Connected if the shared site account is linked, or (legacy) this user is
    acct = _gmail_account_for(u)
    return {
        "connected": bool(acct),
        "site_account": _report_sender_email() if acct == REPORT_SENDER_KEY else "",
    }

@app.post("/api/shifts/fetch-from-gmail")
async def fetch_shifts_from_gmail(body: dict = {}, u: str = Depends(auth)):
    import urllib.request, urllib.parse, urllib.error, json as _json, base64
    acct = _gmail_account_for(u)
    tokens = _get_valid_gmail_tokens(acct) if acct else None
    if not tokens:
        raise HTTPException(400, "Gmail התנתק, יש להתחבר מחדש")
    access_token = tokens["access_token"]

    def gmail_get(url):
        nonlocal access_token
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {access_token}"})
        try:
            with urllib.request.urlopen(req) as r:
                return _json.loads(r.read())
        except urllib.error.HTTPError as e:
            if e.code != 401:
                raise
            # Access token expired mid-flight — refresh once and retry
            refreshed = _refresh_gmail_tokens(acct, tokens)
            if not refreshed:
                raise HTTPException(400, "Gmail התנתק, יש להתחבר מחדש")
            access_token = refreshed["access_token"]
            req = urllib.request.Request(url, headers={"Authorization": f"Bearer {access_token}"})
            with urllib.request.urlopen(req) as r:
                return _json.loads(r.read())

    # Search for email from clock2go, optionally filtered by date in subject
    date_str = body.get("date", "")
    subject_query = f'דו"ח נוכחות כולל משימות יומי {date_str}'.strip()
    q = urllib.parse.quote(f'from:support@clock2go.co.il subject:{subject_query}')
    result = gmail_get(f"https://gmail.googleapis.com/gmail/v1/users/me/messages?q={q}&maxResults=1")
    messages = result.get("messages", [])
    if not messages:
        raise HTTPException(404, "לא נמצא מייל מ-clock2go עם קובץ נוכחות")

    msg = gmail_get(f"https://gmail.googleapis.com/gmail/v1/users/me/messages/{messages[0]['id']}")
    subject = next((h["value"] for h in msg["payload"]["headers"] if h["name"] == "Subject"), "")

    # Find Excel attachment
    def find_parts(part):
        if part.get("filename","").endswith((".xlsx",".xls")) and part.get("body",{}).get("attachmentId"):
            return part
        for p in part.get("parts", []):
            found = find_parts(p)
            if found:
                return found
        return None

    att_part = find_parts(msg["payload"])
    if not att_part:
        raise HTTPException(404, "לא נמצא קובץ Excel במייל")

    att_id = att_part["body"]["attachmentId"]
    att = gmail_get(f"https://gmail.googleapis.com/gmail/v1/users/me/messages/{messages[0]['id']}/attachments/{att_id}")
    file_bytes = base64.urlsafe_b64decode(att["data"])

    # Save to temp and return as upload token
    token = secrets.token_hex(16)
    _excel_cache[token] = file_bytes
    return {"ok": True, "token": token, "filename": att_part["filename"], "subject": subject}


@app.get("/api/recruiter/config")
async def recruiter_config_get(_: str = Depends(auth)):
    return _load_recruiter_config()

@app.post("/api/recruiter/config")
async def recruiter_config_post(body: dict, _: str = Depends(auth)):
    _save_recruiter_config(body)
    return {"ok": True}

@app.post("/api/recruiter/upload")
async def recruiter_upload(excel: UploadFile = File(...), _: str = Depends(auth)):
    import pandas as pd
    from datetime import datetime as _dt

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(await excel.read())
        tmp_path = Path(tmp.name)

    try:
        fname = excel.filename or ""
        if fname.lower().endswith(".csv"):
            df = pd.read_csv(str(tmp_path))
        else:
            df = pd.read_excel(str(tmp_path), sheet_name="פיד", header=0)
    finally:
        tmp_path.unlink(missing_ok=True)

    def _ext(val) -> str | None:
        try:
            s = str(int(float(val)))
            if s.startswith("910") and len(s) >= 6:
                return s[3:].lstrip("0") or s[3:]
            if 2 <= len(s) <= 4:
                return s.lstrip("0") or s
        except Exception:
            pass
        return None

    calls = []
    for _, row in df.iterrows():
        ext = _ext(row.get("src"))
        if not ext:
            continue
        raw_date = row.get("calldate")
        try:
            if pd.isna(raw_date):
                continue
        except Exception:
            if not raw_date:
                continue
        try:
            dt = pd.to_datetime(raw_date)
            date_str = dt.strftime("%Y-%m-%d")
            time_str = dt.strftime("%H:%M")
            hour = int(dt.hour)
        except Exception:
            continue
        try:
            duration = int(float(row.get("billsec") or 0))
        except Exception:
            duration = 0
        answered = str(row.get("disposition", "")).upper() == "ANSWERED"
        dst = str(row.get("dst", "")).rstrip(".0").strip()
        calls.append({
            "date": date_str, "time": time_str, "hour": hour,
            "extension": ext, "dst": dst,
            "duration_sec": duration, "answered": answered,
        })

    _wj(RECRUITER_DATA_FILE, {"last_updated": _dt.now().isoformat(), "calls": calls})
    return {"ok": True, "total_calls": len(calls)}


@app.get("/api/recruiter/data")
async def recruiter_data(
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None),
    _: str = Depends(auth),
):
    import collections

    raw = _rj(RECRUITER_DATA_FILE, None)
    if not raw:
        return {"last_updated": None, "recruiters": [], "all_dates": [], "full_range": None}

    cfg = _load_recruiter_config()
    recruiter_map = {r["extension"]: r["name"] for r in cfg.get("recruiters", [])}
    long_sec = cfg.get("long_call_threshold_minutes", 8) * 60
    repeat_min = cfg.get("repeat_call_threshold", 2)

    all_calls = raw.get("calls", [])
    all_call_dates = sorted(set(c["date"] for c in all_calls))
    full_range = {"from": all_call_dates[0], "to": all_call_dates[-1]} if all_call_dates else None

    df = date_from or (all_call_dates[0] if all_call_dates else "")
    dt = date_to or (all_call_dates[-1] if all_call_dates else "")
    calls = [c for c in all_calls if df <= c["date"] <= dt]
    all_dates = sorted(set(c["date"] for c in calls))

    by_ext: dict = collections.defaultdict(list)
    for c in calls:
        by_ext[c["extension"]].append(c)

    recruiters = []
    for ext, rcalls in by_ext.items():
        name = recruiter_map.get(ext, f"שלוחה {ext}")
        answered = [c for c in rcalls if c["answered"]]
        total = len(rcalls)
        ans_count = len(answered)
        ans_rate = round(ans_count / total * 100, 1) if total else 0
        total_sec = sum(c["duration_sec"] for c in answered)
        total_min = round(total_sec / 60, 1)
        avg_dur = round(total_sec / ans_count / 60, 1) if ans_count else 0
        work_days = len(set(c["date"] for c in rcalls))
        avg_per_day = round(total / work_days, 1) if work_days else 0
        times = sorted(c["time"] for c in rcalls)

        hourly: dict = collections.defaultdict(float)
        for c in answered:
            hourly[c["hour"]] += c["duration_sec"] / 60
        hourly_dist = {str(h): round(hourly.get(h, 0), 1) for h in range(8, 18)}

        daily_calls: dict = collections.defaultdict(int)
        daily_minutes: dict = collections.defaultdict(float)
        for c in rcalls:
            daily_calls[c["date"]] += 1
        for c in answered:
            daily_minutes[c["date"]] += c["duration_sec"] / 60

        long_calls = sorted(
            [{"date": c["date"], "time": c["time"], "dst": c["dst"],
              "minutes": round(c["duration_sec"] / 60, 1)}
             for c in answered if c["duration_sec"] >= long_sec],
            key=lambda x: x["minutes"], reverse=True,
        )

        dst_counts = collections.Counter(c["dst"] for c in rcalls if c["dst"])
        repeat_numbers = [
            {"number": num, "count": cnt}
            for num, cnt in dst_counts.most_common(20)
            if cnt >= repeat_min
        ]

        recruiters.append({
            "name": name, "extension": ext,
            "total_calls": total, "answered_calls": ans_count, "answer_rate": ans_rate,
            "total_minutes": total_min, "avg_duration_minutes": avg_dur,
            "work_days": work_days, "avg_calls_per_day": avg_per_day,
            "first_call": times[0] if times else None,
            "last_call": times[-1] if times else None,
            "hourly_distribution": hourly_dist,
            "daily_calls": dict(daily_calls),
            "daily_minutes": {d: round(v, 1) for d, v in daily_minutes.items()},
            "long_calls": long_calls,
            "repeat_numbers": repeat_numbers,
        })

    known_order = {r["extension"]: i for i, r in enumerate(cfg.get("recruiters", []))}
    recruiters.sort(key=lambda r: known_order.get(r["extension"], 999))

    return {"last_updated": raw.get("last_updated"), "all_dates": all_dates, "full_range": full_range, "recruiters": recruiters}
